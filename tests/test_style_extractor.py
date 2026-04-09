from __future__ import annotations

from io import BytesIO
from pathlib import Path
import sys
import tempfile
import unittest
import zipfile

THIS_DIR = Path(__file__).resolve().parent
SRC_ROOT = THIS_DIR.parent / "src"
if str(SRC_ROOT) not in sys.path:
    sys.path.insert(0, str(SRC_ROOT))

from document_processor.core.style_extractor import extract_styles, extract_styles_docx, extract_styles_hwpx


class StyleExtractorTests(unittest.TestCase):
    def test_extract_docx_paragraph_indents(self) -> None:
        from docx import Document
        from docx.shared import Pt

        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "indent_sample.docx"

            doc = Document()
            paragraph = doc.add_paragraph("Indented")
            paragraph.alignment = 1
            paragraph.paragraph_format.left_indent = Pt(24.0)
            paragraph.paragraph_format.right_indent = Pt(12.0)
            paragraph.paragraph_format.first_line_indent = Pt(-6.0)
            doc.save(str(docx_path))

            style_map = extract_styles_docx(docx_path)
            with docx_path.open("rb") as handle:
                style_map_file = extract_styles(handle.read(), doc_type="docx")

        pstyle = style_map.paragraphs.get("s1.p1")
        self.assertIsNotNone(pstyle)
        assert pstyle is not None

        self.assertEqual(pstyle.align, "center")
        self.assertAlmostEqual(pstyle.left_indent_pt or 0.0, 24.0, places=3)
        self.assertAlmostEqual(pstyle.right_indent_pt or 0.0, 12.0, places=3)
        self.assertAlmostEqual(pstyle.first_line_indent_pt or 0.0, -6.0, places=3)
        self.assertAlmostEqual(style_map_file.paragraphs["s1.p1"].hanging_indent_pt or 0.0, 6.0, places=3)

    def test_extract_hwpx_paragraph_indents(self) -> None:
        header_xml = """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core">
  <hh:paraProperties itemCnt="1">
    <hh:paraPr id="1">
      <hh:align horizontal="CENTER" />
      <hh:margin>
        <hc:intent value="-500" unit="HWPUNIT" />
        <hc:left value="200" unit="HWPUNIT" />
        <hc:right value="300" unit="HWPUNIT" />
      </hh:margin>
    </hh:paraPr>
  </hh:paraProperties>
  <hh:charProperties itemCnt="1">
    <hh:charPr id="1" height="1200" textColor="#112233">
      <hh:bold />
      <hh:underline type="BOTTOM" shape="SOLID" color="#000000" />
      <hh:strikeout shape="NONE" color="#000000" />
    </hh:charPr>
  </hh:charProperties>
</hh:head>
"""

        section_xml = """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec
  xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section"
  xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph"
  xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head"
  xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core">
  <hp:p paraPrIDRef="1">
    <hp:run charPrIDRef="1">
      <hp:t>Hello</hp:t>
    </hp:run>
  </hp:p>
</hs:sec>
"""

        hwpx_bytes_io = BytesIO()
        with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
            zf.writestr("Contents/header.xml", header_xml)
            zf.writestr("Contents/section0.xml", section_xml)
        hwpx_bytes = hwpx_bytes_io.getvalue()

        style_map = extract_styles_hwpx(hwpx_bytes)
        unified = extract_styles(hwpx_bytes, doc_type="hwpx")

        pstyle = style_map.paragraphs.get("s1.p1")
        self.assertIsNotNone(pstyle)
        assert pstyle is not None

        self.assertEqual(pstyle.align, "center")
        self.assertAlmostEqual(pstyle.left_indent_pt or 0.0, 2.0, places=3)
        self.assertAlmostEqual(pstyle.right_indent_pt or 0.0, 3.0, places=3)
        self.assertAlmostEqual(pstyle.first_line_indent_pt or 0.0, -5.0, places=3)
        self.assertAlmostEqual(unified.paragraphs["s1.p1"].hanging_indent_pt or 0.0, 5.0, places=3)

        rstyle = style_map.runs.get("s1.p1.r1")
        self.assertIsNotNone(rstyle)
        assert rstyle is not None
        self.assertTrue(rstyle.bold)
        self.assertTrue(rstyle.underline)
        self.assertEqual(rstyle.color, "#112233")
        self.assertAlmostEqual(rstyle.size_pt or 0.0, 12.0, places=3)

    def test_extract_hwpx_styles_from_hwpx_document(self) -> None:
        from document_processor.hwpx import HwpxDocument

        header_xml = """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core">
  <hh:paraProperties itemCnt="1">
    <hh:paraPr id="1">
      <hh:align horizontal="RIGHT" />
    </hh:paraPr>
  </hh:paraProperties>
</hh:head>
"""
        section_xml = """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p paraPrIDRef="1">
    <hp:run><hp:t>Hello</hp:t></hp:run>
  </hp:p>
</hs:sec>
"""

        hwpx_bytes_io = BytesIO()
        with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
            zf.writestr("Contents/header.xml", header_xml)
            zf.writestr("Contents/section0.xml", section_xml)

        with HwpxDocument.open(hwpx_bytes_io.getvalue()) as doc:
            style_map = extract_styles_hwpx(doc)

        self.assertEqual(style_map.paragraphs["s1.p1"].align, "right")

    def test_extract_hwpx_strikeout_ignores_3d_default_shape(self) -> None:
        header_xml = """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core">
  <hh:charProperties itemCnt="2">
    <hh:charPr id="1" height="1000" textColor="#000000">
      <hh:strikeout shape="3D" color="#000000" />
    </hh:charPr>
    <hh:charPr id="2" height="1000" textColor="#000000">
      <hh:strikeout shape="SOLID" color="#000000" />
    </hh:charPr>
  </hh:charProperties>
</hh:head>
"""

        section_xml = """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec
  xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section"
  xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p>
    <hp:run charPrIDRef="1"><hp:t>A</hp:t></hp:run>
    <hp:run charPrIDRef="2"><hp:t>B</hp:t></hp:run>
  </hp:p>
</hs:sec>
"""

        hwpx_bytes_io = BytesIO()
        with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
            zf.writestr("Contents/header.xml", header_xml)
            zf.writestr("Contents/section0.xml", section_xml)

        style_map = extract_styles_hwpx(hwpx_bytes_io.getvalue())

        self.assertFalse(style_map.runs["s1.p1.r1"].strikethrough)
        self.assertTrue(style_map.runs["s1.p1.r2"].strikethrough)

    def test_extract_hwpx_vertical_merge_uses_logical_cell_ids(self) -> None:
        hwpx_bytes_io = BytesIO()
        with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
            zf.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            zf.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p>
    <hp:run>
      <hp:tbl>
        <hp:tr>
          <hp:tc borderFillIDRef="1"><hp:subList vertAlign="CENTER"><hp:p><hp:run><hp:t>Main</hp:t></hp:run></hp:p></hp:subList><hp:cellAddr colAddr="0" rowAddr="0"/><hp:cellSpan colSpan="1" rowSpan="4"/></hp:tc>
          <hp:tc borderFillIDRef="1"><hp:subList vertAlign="CENTER"><hp:p><hp:run><hp:t>관</hp:t></hp:run></hp:p></hp:subList><hp:cellAddr colAddr="1" rowAddr="0"/><hp:cellSpan colSpan="1" rowSpan="1"/></hp:tc>
        </hp:tr>
        <hp:tr>
          <hp:tc borderFillIDRef="1"><hp:subList vertAlign="CENTER"><hp:p><hp:run><hp:t>항</hp:t></hp:run></hp:p></hp:subList><hp:cellAddr colAddr="1" rowAddr="1"/><hp:cellSpan colSpan="1" rowSpan="1"/></hp:tc>
        </hp:tr>
      </hp:tbl>
    </hp:run>
  </hp:p>
</hs:sec>
""",
            )
        style_map = extract_styles_hwpx(hwpx_bytes_io.getvalue())

        self.assertIn("s1.p1.r1.tbl1.tr1.tc2", style_map.cells)
        self.assertIn("s1.p1.r1.tbl1.tr2.tc2", style_map.cells)
        self.assertNotIn("s1.p1.r1.tbl1.tr2.tc1", style_map.cells)
        self.assertEqual(style_map.tables["s1.p1.r1.tbl1"].row_count, 4)
        self.assertEqual(style_map.tables["s1.p1.r1.tbl1"].col_count, 2)

    def test_extract_docx_nested_table_styles(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "nested_styles.docx"
            doc = Document()
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            nested = cell.add_table(rows=1, cols=1)
            nested.cell(0, 0).paragraphs[0].alignment = 1
            doc.save(str(docx_path))

            style_map = extract_styles_docx(docx_path)

        self.assertIn("s1.p1.r1.tbl1.tr1.tc1.p1.tbl1", style_map.tables)
        self.assertIn("s1.p1.r1.tbl1.tr1.tc1.p1.tbl1.tr1.tc1.p1", style_map.paragraphs)

    def test_extract_docx_table_grid_borders_from_table_style(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "table_grid.docx"
            doc = Document()
            table = doc.add_table(rows=2, cols=2)
            table.style = "Table Grid"
            table.cell(0, 0).text = "A1"
            table.cell(0, 1).text = "A2"
            table.cell(1, 0).text = "B1"
            table.cell(1, 1).text = "B2"
            doc.save(str(docx_path))

            style_map = extract_styles_docx(docx_path)

        top_left = style_map.cells["s1.p1.r1.tbl1.tr1.tc1"]
        bottom_right = style_map.cells["s1.p1.r1.tbl1.tr2.tc2"]

        self.assertEqual(top_left.border_top, "1px solid #000000")
        self.assertEqual(top_left.border_left, "1px solid #000000")
        self.assertEqual(top_left.border_right, "1px solid #000000")
        self.assertEqual(top_left.border_bottom, "1px solid #000000")
        self.assertEqual(bottom_right.border_top, "1px solid #000000")
        self.assertEqual(bottom_right.border_left, "1px solid #000000")
        self.assertEqual(bottom_right.border_right, "1px solid #000000")
        self.assertEqual(bottom_right.border_bottom, "1px solid #000000")

    def test_extract_docx_table_and_cell_sizes(self) -> None:
        from docx import Document
        from docx.shared import Inches, Pt

        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "sized_table.docx"
            doc = Document()
            table = doc.add_table(rows=2, cols=2)
            table.style = "Table Grid"
            table.autofit = False
            table.columns[0].width = Inches(1.0)
            table.columns[1].width = Inches(2.0)
            for row in table.rows:
                row.cells[0].width = Inches(1.0)
                row.cells[1].width = Inches(2.0)
                row.height = Pt(24.0)
            doc.save(str(docx_path))

            style_map = extract_styles_docx(docx_path)

        table_style = style_map.tables["s1.p1.r1.tbl1"]
        first_cell = style_map.cells["s1.p1.r1.tbl1.tr1.tc1"]
        second_cell = style_map.cells["s1.p1.r1.tbl1.tr1.tc2"]

        self.assertAlmostEqual(table_style.width_pt or 0.0, 216.0, places=1)
        self.assertAlmostEqual(first_cell.width_pt or 0.0, 72.0, places=1)
        self.assertAlmostEqual(second_cell.width_pt or 0.0, 144.0, places=1)
        self.assertAlmostEqual(first_cell.height_pt or 0.0, 24.0, places=1)

    def test_extract_docx_cell_diagonal_borders(self) -> None:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "diag.docx"
            doc = Document()
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            cell.text = "Diag"

            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)

            for side, value, color in (
                ("w:tl2br", "single", "000000"),
                ("w:tr2bl", "dashed", "FF0000"),
            ):
                border = OxmlElement(side)
                border.set(qn("w:val"), value)
                border.set(qn("w:sz"), "4")
                border.set(qn("w:color"), color)
                tc_borders.append(border)

            doc.save(str(docx_path))

            style_map = extract_styles_docx(docx_path)

        cell_style = style_map.cells["s1.p1.r1.tbl1.tr1.tc1"]
        self.assertEqual(cell_style.diagonal_tl_br, "1px solid #000000")
        self.assertEqual(cell_style.diagonal_tr_bl, "1px dashed #FF0000")

    def test_extract_hwpx_cell_diagonal_borders(self) -> None:
        header_xml = """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core">
  <hh:borderFills itemCnt="2">
    <hh:borderFill id="1" threeD="0" shadow="0" centerLine="NONE" breakCellSeparateLine="0">
      <hh:slash type="CENTER" Crooked="0" isCounter="0" />
      <hh:backSlash type="NONE" Crooked="0" isCounter="0" />
      <hh:leftBorder type="SOLID" width="0.12 mm" color="#000000" />
      <hh:rightBorder type="SOLID" width="0.12 mm" color="#000000" />
      <hh:topBorder type="SOLID" width="0.12 mm" color="#000000" />
      <hh:bottomBorder type="SOLID" width="0.12 mm" color="#000000" />
      <hh:diagonal type="SOLID" width="0.12 mm" color="#123456" />
    </hh:borderFill>
    <hh:borderFill id="2" threeD="0" shadow="0" centerLine="NONE" breakCellSeparateLine="0">
      <hh:slash type="NONE" Crooked="0" isCounter="0" />
      <hh:backSlash type="CENTER" Crooked="0" isCounter="0" />
      <hh:leftBorder type="SOLID" width="0.12 mm" color="#000000" />
      <hh:rightBorder type="SOLID" width="0.12 mm" color="#000000" />
      <hh:topBorder type="SOLID" width="0.12 mm" color="#000000" />
      <hh:bottomBorder type="SOLID" width="0.12 mm" color="#000000" />
      <hh:diagonal type="DASH" width="0.12 mm" color="#654321" />
    </hh:borderFill>
  </hh:borderFills>
</hh:head>
"""

        section_xml = """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec
  xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section"
  xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p>
    <hp:run>
      <hp:tbl>
        <hp:tr>
          <hp:tc borderFillIDRef="1">
            <hp:subList vertAlign="CENTER">
              <hp:p><hp:run><hp:t>A</hp:t></hp:run></hp:p>
            </hp:subList>
            <hp:cellAddr colAddr="0" rowAddr="0" />
            <hp:cellSpan colSpan="1" rowSpan="1" />
          </hp:tc>
          <hp:tc borderFillIDRef="2">
            <hp:subList vertAlign="CENTER">
              <hp:p><hp:run><hp:t>B</hp:t></hp:run></hp:p>
            </hp:subList>
            <hp:cellAddr colAddr="1" rowAddr="0" />
            <hp:cellSpan colSpan="1" rowSpan="1" />
          </hp:tc>
        </hp:tr>
      </hp:tbl>
    </hp:run>
  </hp:p>
</hs:sec>
"""

        hwpx_bytes_io = BytesIO()
        with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
            zf.writestr("Contents/header.xml", header_xml)
            zf.writestr("Contents/section0.xml", section_xml)

        style_map = extract_styles_hwpx(hwpx_bytes_io.getvalue())

        slash_cell = style_map.cells["s1.p1.r1.tbl1.tr1.tc1"]
        backslash_cell = style_map.cells["s1.p1.r1.tbl1.tr1.tc2"]
        self.assertEqual(slash_cell.diagonal_tr_bl, "1px solid #123456")
        self.assertIsNone(slash_cell.diagonal_tl_br)
        self.assertEqual(backslash_cell.diagonal_tl_br, "1px dashed #654321")
        self.assertIsNone(backslash_cell.diagonal_tr_bl)

    def test_extract_hwpx_nested_table_styles(self) -> None:
        hwpx_bytes_io = BytesIO()
        with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
            zf.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core">
  <hh:paraProperties itemCnt="1"><hh:paraPr id="1"><hh:align horizontal="CENTER" /></hh:paraPr></hh:paraProperties>
  <hh:charProperties itemCnt="1"><hh:charPr id="1" height="1200" /></hh:charProperties>
</hh:head>
""",
            )
            zf.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p paraPrIDRef="1">
    <hp:run charPrIDRef="1">
      <hp:tbl>
        <hp:tr>
          <hp:tc>
            <hp:subList>
              <hp:p paraPrIDRef="1">
                <hp:run charPrIDRef="1"><hp:t>Outer</hp:t></hp:run>
                <hp:run charPrIDRef="1">
                  <hp:tbl>
                    <hp:tr>
                      <hp:tc>
                        <hp:subList>
                          <hp:p paraPrIDRef="1"><hp:run charPrIDRef="1"><hp:t>Inner</hp:t></hp:run></hp:p>
                        </hp:subList>
                        <hp:cellAddr colAddr="0" rowAddr="0"/>
                        <hp:cellSpan colSpan="1" rowSpan="1"/>
                      </hp:tc>
                    </hp:tr>
                  </hp:tbl>
                </hp:run>
              </hp:p>
            </hp:subList>
            <hp:cellAddr colAddr="0" rowAddr="0"/>
            <hp:cellSpan colSpan="1" rowSpan="1"/>
          </hp:tc>
        </hp:tr>
      </hp:tbl>
    </hp:run>
  </hp:p>
</hs:sec>
""",
            )
        style_map = extract_styles_hwpx(hwpx_bytes_io.getvalue())

        self.assertIn("s1.p1.r1.tbl1.tr1.tc1.p1.tbl1", style_map.tables)
        self.assertIn("s1.p1.r1.tbl1.tr1.tc1.p1.tbl1.tr1.tc1.p1", style_map.paragraphs)

    def test_extract_hwpx_table_and_cell_sizes(self) -> None:
        hwpx_bytes_io = BytesIO()
        with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
            zf.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            zf.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p>
    <hp:run>
      <hp:tbl rowCnt="1" colCnt="1">
        <hp:sz width="14400" height="4800" widthRelTo="ABSOLUTE" heightRelTo="ABSOLUTE" protect="0"/>
        <hp:tr>
          <hp:tc>
            <hp:subList><hp:p><hp:run><hp:t>Cell</hp:t></hp:run></hp:p></hp:subList>
            <hp:cellAddr colAddr="0" rowAddr="0"/>
            <hp:cellSpan colSpan="1" rowSpan="1"/>
            <hp:cellSz width="7200" height="3600"/>
          </hp:tc>
        </hp:tr>
      </hp:tbl>
    </hp:run>
  </hp:p>
</hs:sec>
""",
            )

        style_map = extract_styles_hwpx(hwpx_bytes_io.getvalue())

        table_style = style_map.tables["s1.p1.r1.tbl1"]
        cell_style = style_map.cells["s1.p1.r1.tbl1.tr1.tc1"]

        self.assertAlmostEqual(table_style.width_pt or 0.0, 144.0, places=3)
        self.assertAlmostEqual(table_style.height_pt or 0.0, 48.0, places=3)
        self.assertAlmostEqual(cell_style.width_pt or 0.0, 72.0, places=3)
        self.assertAlmostEqual(cell_style.height_pt or 0.0, 36.0, places=3)


if __name__ == "__main__":
    unittest.main()
