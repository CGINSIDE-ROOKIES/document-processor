from __future__ import annotations

import base64
from io import BytesIO
from pathlib import Path
import sys
import tempfile
import unittest
import zipfile

from pydantic import BaseModel

THIS_DIR = Path(__file__).resolve().parent
SRC_ROOT = THIS_DIR.parent / "src"
if str(SRC_ROOT) not in sys.path:
    sys.path.insert(0, str(SRC_ROOT))

from document_processor import DocIR, ImageAsset, ImageIR, ParagraphIR, RunIR


PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO2Z4vsAAAAASUVORK5CYII="
)


class ImageSupportTests(unittest.TestCase):
    def test_image_metadata_lives_on_shared_asset(self) -> None:
        class ImageMeta(BaseModel):
            label: str

        asset = ImageAsset[ImageMeta](
            mime_type="image/png",
            filename="shared.png",
            data_base64="AAAA",
            meta=ImageMeta(label="shared"),
        )
        doc = DocIR(
            assets={"img1": asset},
            paragraphs=[
                ParagraphIR(
                    unit_id="s1.p1",
                    content=[
                        ImageIR(unit_id="s1.p1.img1", image_id="img1", display_width_pt=50.0),
                        ImageIR(unit_id="s1.p1.img2", image_id="img1", display_width_pt=90.0),
                    ],
                )
            ],
        )

        first_image, second_image = doc.paragraphs[0].images
        self.assertFalse(hasattr(first_image, "meta"))
        self.assertIs(doc.get_image_asset(first_image), asset)
        self.assertIs(doc.get_image_asset(second_image), asset)
        self.assertEqual(doc.get_image_asset("img1").meta.label, "shared")
        self.assertEqual(first_image.display_width_pt, 50.0)
        self.assertEqual(second_image.display_width_pt, 90.0)

        dumped = doc.model_dump(mode="json")
        self.assertEqual(dumped["assets"]["img1"]["meta"]["label"], "shared")
        self.assertNotIn("meta", dumped["paragraphs"][0]["content"][0])

    def test_docx_images_are_embedded_in_ir_and_preserve_order(self) -> None:
        from docx import Document
        from docx.shared import Inches

        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "images.docx"
            doc = Document()
            paragraph = doc.add_paragraph()
            paragraph.add_run("Before ")
            paragraph.add_run().add_picture(BytesIO(PNG_BYTES), width=Inches(1.0))
            paragraph.add_run("After")
            doc.save(str(docx_path))

            parsed = DocIR.from_file(docx_path)

        paragraph_ir = next(paragraph for paragraph in parsed.paragraphs if paragraph.content)
        self.assertEqual(
            [type(node).__name__ for node in paragraph_ir.content],
            ["RunIR", "ImageIR", "RunIR"],
        )
        self.assertEqual(paragraph_ir.runs[0].text, "Before ")
        self.assertEqual(paragraph_ir.runs[1].text, "After")
        self.assertEqual(len(parsed.assets), 1)

        image = paragraph_ir.images[0]
        asset = parsed.assets[image.image_id]
        self.assertEqual(asset.mime_type, "image/png")
        self.assertTrue(asset.data_base64)
        self.assertEqual(asset.intrinsic_width_px, 1)
        self.assertEqual(asset.intrinsic_height_px, 1)
        self.assertAlmostEqual(image.display_width_pt or 0.0, 72.0, places=1)

        dumped = parsed.model_dump(mode="json")
        self.assertIn(image.image_id, dumped["assets"])
        self.assertEqual(dumped["assets"][image.image_id]["mime_type"], "image/png")

    def test_hwpx_images_are_embedded_in_ir_and_render_in_html(self) -> None:
        hwpx_bytes = self._build_hwpx_with_inline_image()

        parsed = DocIR.from_file(hwpx_bytes, doc_type="hwpx")
        paragraph_ir = parsed.paragraphs[0]

        self.assertEqual(
            [type(node).__name__ for node in paragraph_ir.content],
            ["RunIR", "ImageIR", "RunIR"],
        )
        self.assertIsInstance(paragraph_ir.content[1], ImageIR)
        self.assertEqual(len(parsed.assets), 1)
        self.assertEqual(parsed.assets[paragraph_ir.images[0].image_id].intrinsic_width_px, 1)
        self.assertEqual(parsed.assets[paragraph_ir.images[0].image_id].intrinsic_height_px, 1)
        self.assertAlmostEqual(paragraph_ir.images[0].display_width_pt or 0.0, 72.0, places=3)
        self.assertAlmostEqual(paragraph_ir.images[0].display_height_pt or 0.0, 36.0, places=3)

        html = parsed.to_html(title="Images")
        self.assertIn("<img ", html)
        self.assertIn("data:image/png;base64,", html)
        self.assertIn("width:72.0pt", html)
        self.assertIn("height:36.0pt", html)
        self.assertIn("Before", html)
        self.assertIn("After", html)

    @staticmethod
    def _build_hwpx_with_inline_image() -> bytes:
        hwpx_bytes_io = BytesIO()
        with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
            zf.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            zf.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section"
        xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph"
        xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core">
  <hp:p>
    <hp:run><hp:t>Before</hp:t></hp:run>
    <hp:run>
      <hp:pic>
        <hp:imgDim dimwidth="7200" dimheight="3600"/>
        <hc:img binaryItemIDRef="image1"/>
      </hp:pic>
    </hp:run>
    <hp:run><hp:t>After</hp:t></hp:run>
  </hp:p>
</hs:sec>
""",
            )
            zf.writestr("BinData/image1.png", PNG_BYTES)
        return hwpx_bytes_io.getvalue()


if __name__ == "__main__":
    unittest.main()
