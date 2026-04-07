from __future__ import annotations

from io import BytesIO
from pathlib import Path
import sys
import tempfile
import unittest
from unittest.mock import patch
import zipfile

THIS_DIR = Path(__file__).resolve().parent
SRC_ROOT = THIS_DIR.parent / "src"
if str(SRC_ROOT) not in sys.path:
    sys.path.insert(0, str(SRC_ROOT))

from document_processor import (
    CellStyleInfo,
    DocIR,
    ImageIR,
    ParaStyleInfo,
    ParagraphIR,
    RunIR,
    RunStyleInfo,
    StyleMap,
    TableIR,
    TableStyleInfo,
    build_doc_ir_from_mapping,
)
from document_processor.core.hwpx_structured_exporter import export_hwpx_structured_mapping


class DocumentIRTests(unittest.TestCase):
    def _sample_mapping(self) -> dict[str, str]:
        return {
            "s1.p1.r1": "Hello ",
            "s1.p1.r2": "World",
            "s1.p2.r1.tbl1.tr1.tc1.p1.r1": "A1",
            "s1.p2.r1.tbl1.tr1.tc2.p1.r1": "B1",
            "s1.p2.r1.tbl1.tr2.tc1.p1.r1": "A2",
            "s1.p2.r1.tbl1.tr2.tc2.p1.r1": "B2",
        }

    def _sample_style_map(self) -> StyleMap:
        return StyleMap(
            runs={
                "s1.p1.r1": RunStyleInfo(bold=True, size_pt=11.0),
                "s1.p1.r2": RunStyleInfo(italic=True, size_pt=11.0),
            },
            paragraphs={
                "s1.p1": ParaStyleInfo(align="center"),
            },
            cells={
                "s1.p2.r1.tbl1.tr1.tc1": CellStyleInfo(background="#ffeeaa"),
            },
            tables={
                "s1.p2.r1.tbl1": TableStyleInfo(row_count=2, col_count=2),
            },
        )

    def test_hierarchy_construction(self) -> None:
        doc_ir = build_doc_ir_from_mapping(self._sample_mapping())

        self.assertEqual(len(doc_ir.paragraphs), 2)
        self.assertEqual(doc_ir.paragraphs[0].text, "Hello World")
        self.assertEqual(doc_ir.paragraphs[1].content[0].unit_id, "s1.p2.r1.tbl1")
        self.assertEqual(doc_ir.paragraphs[1].tables[0].row_count, 2)
        self.assertEqual(doc_ir.paragraphs[1].tables[0].col_count, 2)

    def test_style_embedding(self) -> None:
        doc_ir = build_doc_ir_from_mapping(self._sample_mapping(), style_map=self._sample_style_map())
        self.assertEqual(doc_ir.paragraphs[0].para_style.align, "center")
        self.assertTrue(doc_ir.paragraphs[0].runs[0].run_style.bold)
        self.assertEqual(doc_ir.paragraphs[1].tables[0].cells[0].cell_style.background, "#ffeeaa")

    def test_docir_subclass_from_mapping(self) -> None:
        class DocumentLM(DocIR):
            custom_field: int = 0

        doc = DocumentLM.from_mapping({"s1.p1.r1": "X"}, custom_field=7)
        self.assertIsInstance(doc, DocumentLM)
        self.assertEqual(doc.custom_field, 7)

    def test_content_is_source_of_truth(self) -> None:
        doc = DocIR(
            paragraphs=[
                ParagraphIR(
                    unit_id="s1.p1",
                    content=[
                        RunIR(unit_id="s1.p1.r1", text="Hello"),
                        ImageIR(unit_id="s1.p1.img1", image_id="img1"),
                        TableIR(unit_id="s1.p1.tbl1"),
                    ],
                )
            ]
        )

        paragraph = doc.paragraphs[0]
        self.assertEqual(
            [type(node).__name__ for node in paragraph.content],
            ["RunIR", "ImageIR", "TableIR"],
        )
        self.assertEqual([run.text for run in paragraph.runs], ["Hello"])
        self.assertEqual([image.image_id for image in paragraph.images], ["img1"])
        self.assertEqual([table.unit_id for table in paragraph.tables], ["s1.p1.tbl1"])

        content_annotation = ParagraphIR.model_fields["content"].annotation
        self.assertIn("RunIR", str(content_annotation))
        self.assertIn("ImageIR", str(content_annotation))
        self.assertIn("TableIR", str(content_annotation))

    def test_from_file_docx_path_and_file_object(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "sample.docx"

            doc = Document()
            doc.add_paragraph("Hello")
            doc.save(str(docx_path))

            from_path = DocIR.from_file(docx_path)
            with docx_path.open("rb") as handle:
                from_file_object = DocIR.from_file(handle)

        self.assertEqual(from_path.source_doc_type, "docx")
        self.assertEqual(from_path.source_path, str(docx_path))
        self.assertEqual(from_path.paragraphs[0].text, "Hello")
        self.assertEqual(from_file_object.source_doc_type, "docx")
        self.assertEqual(from_file_object.paragraphs[0].text, "Hello")

    def test_from_file_hwpx_bytes_and_file_object(self) -> None:
        hwpx_bytes_io = BytesIO()
        with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
            zf.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            zf.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p><hp:run><hp:t>Hello HWPX</hp:t></hp:run></hp:p>
</hs:sec>
""",
            )
        hwpx_bytes = hwpx_bytes_io.getvalue()

        from_bytes = DocIR.from_file(hwpx_bytes, doc_type="hwpx")
        from_file_object = DocIR.from_file(BytesIO(hwpx_bytes), doc_type="hwpx")

        self.assertEqual(from_bytes.source_doc_type, "hwpx")
        self.assertEqual(from_bytes.paragraphs[0].text, "Hello HWPX")
        self.assertEqual(from_file_object.paragraphs[0].text, "Hello HWPX")

    def test_from_file_hwpx_reads_mixed_content_inside_hp_t(self) -> None:
        hwpx_bytes_io = BytesIO()
        with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
            zf.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            zf.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p>
    <hp:run>
      <hp:t>&lt;<hp:fwSpace />수요기업 협업 규모 및 분야<hp:fwSpace />&gt;</hp:t>
    </hp:run>
  </hp:p>
</hs:sec>
""",
            )

        doc = DocIR.from_file(hwpx_bytes_io.getvalue(), doc_type="hwpx")

        self.assertEqual(doc.paragraphs[0].text, "<수요기업 협업 규모 및 분야>")

    def test_from_file_hwp_file_object_materializes_temp_path(self) -> None:
        fake_hwp = b"fake-hwp"

        with (
            patch("document_processor.core.document_ir_parser.convert_hwp_to_hwpx_bytes") as convert_hwp,
            patch("document_processor.core.style_extractor.extract_styles") as extract_styles,
        ):
            hwpx_bytes_io = BytesIO()
            with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
                zf.writestr(
                    "Contents/header.xml",
                    """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
                )
                zf.writestr(
                    "Contents/section0.xml",
                    """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p><hp:run><hp:t>Converted</hp:t></hp:run></hp:p>
</hs:sec>
""",
                )
            convert_hwp.return_value = hwpx_bytes_io.getvalue()
            extract_styles.return_value = StyleMap()

            doc = DocIR.from_file(BytesIO(fake_hwp), doc_type="hwp")

        self.assertEqual(doc.source_doc_type, "hwp")
        self.assertEqual(doc.paragraphs[0].text, "Converted")
        convert_source = convert_hwp.call_args.kwargs.get("hwp_path", convert_hwp.call_args.args[0])
        self.assertTrue(isinstance(convert_source, Path))
        self.assertEqual(convert_source.suffix, ".hwp")

    def test_hwpx_vertical_merge_uses_logical_column_ids(self) -> None:
        hwpx_bytes_io = BytesIO()
        with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
            zf.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            zf.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p>
    <hp:run>
      <hp:tbl>
        <hp:tr>
          <hp:tc><hp:subList><hp:p><hp:run><hp:t>Main</hp:t></hp:run></hp:p></hp:subList><hp:cellAddr colAddr="0" rowAddr="0"/><hp:cellSpan colSpan="1" rowSpan="4"/></hp:tc>
          <hp:tc><hp:subList><hp:p><hp:run><hp:t>관</hp:t></hp:run></hp:p></hp:subList><hp:cellAddr colAddr="1" rowAddr="0"/><hp:cellSpan colSpan="1" rowSpan="1"/></hp:tc>
          <hp:tc><hp:subList><hp:p><hp:run><hp:t>A</hp:t></hp:run></hp:p></hp:subList><hp:cellAddr colAddr="2" rowAddr="0"/><hp:cellSpan colSpan="1" rowSpan="1"/></hp:tc>
        </hp:tr>
        <hp:tr>
          <hp:tc><hp:subList><hp:p><hp:run><hp:t>항</hp:t></hp:run></hp:p></hp:subList><hp:cellAddr colAddr="1" rowAddr="1"/><hp:cellSpan colSpan="1" rowSpan="1"/></hp:tc>
          <hp:tc><hp:subList><hp:p><hp:run><hp:t>B</hp:t></hp:run></hp:p></hp:subList><hp:cellAddr colAddr="2" rowAddr="1"/><hp:cellSpan colSpan="1" rowSpan="1"/></hp:tc>
        </hp:tr>
      </hp:tbl>
    </hp:run>
  </hp:p>
</hs:sec>
""",
            )
        mapping = export_hwpx_structured_mapping(hwpx_bytes_io.getvalue())

        self.assertIn("s1.p1.r1.tbl1.tr1.tc2.p1.r1", mapping)
        self.assertIn("s1.p1.r1.tbl1.tr2.tc2.p1.r1", mapping)
        self.assertIn("s1.p1.r1.tbl1.tr2.tc3.p1.r1", mapping)
        self.assertNotIn("s1.p1.r1.tbl1.tr2.tc1.p1.r1", mapping)

    def test_builder_supports_nested_tables_in_cell_paragraphs(self) -> None:
        mapping = {
            "s1.p1.r1.tbl1.tr1.tc1.p1.r1": "Outer",
            "s1.p1.r1.tbl1.tr1.tc1.p1.tbl1.tr1.tc1.p1.r1": "Inner",
        }

        doc = DocIR.from_mapping(mapping)
        outer_cell_paragraph = doc.paragraphs[0].tables[0].cells[0].paragraphs[0]

        self.assertEqual(outer_cell_paragraph.text, "Outer\nInner")
        self.assertEqual(len(outer_cell_paragraph.tables), 1)
        self.assertEqual(outer_cell_paragraph.tables[0].unit_id, "s1.p1.r1.tbl1.tr1.tc1.p1.tbl1")
        self.assertEqual(
            outer_cell_paragraph.tables[0].cells[0].paragraphs[0].runs[0].text,
            "Inner",
        )

    def test_docx_nested_tables_are_parsed(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "nested.docx"
            doc = Document()
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            cell.text = "Outer"
            nested = cell.add_table(rows=1, cols=1)
            nested.cell(0, 0).text = "Inner"
            doc.save(str(docx_path))

            parsed = DocIR.from_file(docx_path)

        outer_cell_paragraph = parsed.paragraphs[0].tables[0].cells[0].paragraphs[0]
        self.assertEqual(outer_cell_paragraph.runs[0].text, "Outer")
        self.assertEqual(len(outer_cell_paragraph.tables), 1)
        self.assertEqual(
            outer_cell_paragraph.tables[0].cells[0].paragraphs[0].runs[0].text,
            "Inner",
        )

    def test_hwpx_nested_tables_are_parsed(self) -> None:
        hwpx_bytes_io = BytesIO()
        with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
            zf.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            zf.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p>
    <hp:run>
      <hp:tbl>
        <hp:tr>
          <hp:tc>
            <hp:subList>
              <hp:p>
                <hp:run><hp:t>Outer</hp:t></hp:run>
                <hp:run>
                  <hp:tbl>
                    <hp:tr>
                      <hp:tc>
                        <hp:subList>
                          <hp:p><hp:run><hp:t>Inner</hp:t></hp:run></hp:p>
                        </hp:subList>
                        <hp:cellAddr colAddr="0" rowAddr="0"/>
                        <hp:cellSpan colSpan="1" rowSpan="1"/>
                      </hp:tc>
                    </hp:tr>
                  </hp:tbl>
                </hp:run>
              </hp:p>
            </hp:subList>
            <hp:cellAddr colAddr="0" rowAddr="0"/>
            <hp:cellSpan colSpan="1" rowSpan="1"/>
          </hp:tc>
        </hp:tr>
      </hp:tbl>
    </hp:run>
  </hp:p>
</hs:sec>
""",
            )

        parsed = DocIR.from_file(hwpx_bytes_io.getvalue(), doc_type="hwpx")

        outer_cell_paragraph = parsed.paragraphs[0].tables[0].cells[0].paragraphs[0]
        self.assertEqual(outer_cell_paragraph.runs[0].text, "Outer")
        self.assertEqual(len(outer_cell_paragraph.tables), 1)
        self.assertEqual(
            outer_cell_paragraph.tables[0].cells[0].paragraphs[0].runs[0].text,
            "Inner",
        )


if __name__ == "__main__":
    unittest.main()
