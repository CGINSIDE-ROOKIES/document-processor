from __future__ import annotations

from io import BytesIO
from pathlib import Path
import tempfile
import unittest
import zipfile

from document_processor import (
    ApplyTextEditsRequest,
    DocumentInput,
    DocIR,
    GetDocumentContextRequest,
    ListEditableTargetsRequest,
    ReadDocumentRequest,
    RenderReviewHtmlRequest,
    TextAnnotation,
    TextEdit,
    apply_text_edits,
    get_document_context,
    list_editable_targets,
    read_document,
    render_review_html,
)


class EditorApiTests(unittest.TestCase):
    @staticmethod
    def _build_sample_docx_bytes() -> bytes:
        from docx import Document

        docx = Document()
        paragraph = docx.add_paragraph()
        paragraph.add_run("Hello ")
        paragraph.add_run("World")
        docx.add_paragraph("Second paragraph")

        buffer = BytesIO()
        docx.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _build_sample_table_docx_bytes() -> bytes:
        from docx import Document

        docx = Document()
        table = docx.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Left"
        table.cell(0, 1).text = "Right"

        buffer = BytesIO()
        docx.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _build_sample_hwpx_bytes() -> bytes:
        hwpx_bytes = BytesIO()
        with zipfile.ZipFile(hwpx_bytes, "w") as archive:
            archive.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            archive.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p>
    <hp:run><hp:t>Hello </hp:t></hp:run>
    <hp:run><hp:t>World</hp:t></hp:run>
  </hp:p>
</hs:sec>
""",
            )
        return hwpx_bytes.getvalue()

    @staticmethod
    def _build_sample_table_hwpx_bytes() -> bytes:
        hwpx_bytes = BytesIO()
        with zipfile.ZipFile(hwpx_bytes, "w") as archive:
            archive.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            archive.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p>
    <hp:run>
      <hp:tbl>
        <hp:tr>
          <hp:tc>
            <hp:subList>
              <hp:p><hp:run><hp:t>Left</hp:t></hp:run></hp:p>
            </hp:subList>
            <hp:cellAddr colAddr="0" rowAddr="0"/>
            <hp:cellSpan colSpan="1" rowSpan="1"/>
          </hp:tc>
          <hp:tc>
            <hp:subList>
              <hp:p><hp:run><hp:t>Right</hp:t></hp:run></hp:p>
            </hp:subList>
            <hp:cellAddr colAddr="1" rowAddr="0"/>
            <hp:cellSpan colSpan="1" rowSpan="1"/>
          </hp:tc>
        </hp:tr>
      </hp:tbl>
    </hp:run>
  </hp:p>
</hs:sec>
""",
            )
        return hwpx_bytes.getvalue()

    @staticmethod
    def _build_namespaced_hwpx_bytes() -> bytes:
        hwpx_bytes = BytesIO()
        with zipfile.ZipFile(hwpx_bytes, "w") as archive:
            archive.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            archive.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8" standalone="yes" ?><hs:sec xmlns:ha="http://www.hancom.co.kr/hwpml/2011/app" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph" xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core"><hc:pt0 x="0" y="0"/><hp:p><hp:run><hp:t>Hello </hp:t></hp:run><hp:run><hp:t>World</hp:t></hp:run></hp:p></hs:sec>
""",
            )
        return hwpx_bytes.getvalue()

    def test_get_document_context_accepts_bytes_backed_input(self) -> None:
        source_bytes = self._build_sample_docx_bytes()
        doc = DocIR.from_file(source_bytes, doc_type="docx")
        target_id = doc.paragraphs[0].runs[1].node_id

        result = get_document_context(
            GetDocumentContextRequest(
                document=DocumentInput(
                    source_bytes=source_bytes,
                    source_name="sample.docx",
                ),
                target_ids=[target_id],
                before=0,
                after=1,
            )
        )

        self.assertEqual(result.source_name, "sample.docx")
        self.assertEqual([paragraph.node_id for paragraph in result.paragraphs], [doc.paragraphs[0].node_id, doc.paragraphs[1].node_id])
        self.assertTrue(result.paragraphs[0].node_id.startswith("p_"))
        self.assertEqual(result.paragraphs[0].runs[1].text, "World")
        self.assertTrue(result.paragraphs[0].runs[1].node_id.startswith("r_"))
        self.assertEqual(result.paragraphs[0].text, "Hello World")
        self.assertEqual(
            [(run.text, run.start, run.end) for run in result.paragraphs[0].runs],
            [("Hello ", 0, 6), ("World", 6, 11)],
        )

    def test_read_document_returns_bounded_stable_ids(self) -> None:
        doc = DocIR.from_mapping(
            {
                "s1.p1.r1": "First",
                "s1.p2.r1": "Second",
                "s1.p3.r1": "Third",
            },
            source_doc_type="docx",
        )

        result = read_document(
            ReadDocumentRequest(
                document=DocumentInput(doc_ir=doc),
                start=1,
                limit=1,
            )
        )

        self.assertEqual(result.total_paragraphs, 3)
        self.assertEqual(result.next_start, 2)
        self.assertEqual(result.paragraphs[0].text, "Second")
        self.assertEqual(result.paragraphs[0].node_id, doc.paragraphs[1].node_id)
        self.assertEqual(result.paragraphs[0].native_anchor.debug_path, "s1.p2")
        self.assertEqual(result.paragraphs[0].runs[0].start, 0)
        self.assertEqual(result.paragraphs[0].runs[0].end, len("Second"))

    def test_apply_text_edits_returns_output_bytes_for_bytes_backed_source(self) -> None:
        source_bytes = self._build_sample_docx_bytes()
        doc = DocIR.from_file(source_bytes, doc_type="docx")

        result = apply_text_edits(
            ApplyTextEditsRequest(
                document=DocumentInput(
                    source_bytes=source_bytes,
                    source_name="sample.docx",
                ),
                edits=[
                    TextEdit(
                        target_kind="paragraph",
                        target_id=doc.paragraphs[0].node_id,
                        expected_text="Hello World",
                        new_text="Hello Legal World",
                        reason="Expand wording",
                    )
                ],
                return_doc_ir=True,
            )
        )

        self.assertTrue(result.ok)
        self.assertEqual(result.output_filename, "sample_edited.docx")
        self.assertIsNone(result.output_path)
        self.assertIsNotNone(result.output_bytes)
        self.assertIsNotNone(result.updated_doc_ir)
        self.assertEqual(DocIR.from_file(result.output_bytes).paragraphs[0].text, "Hello Legal World")
        self.assertEqual(result.updated_doc_ir.paragraphs[0].text, "Hello Legal World")

    def test_apply_text_edits_with_doc_ir_only_returns_updated_doc_ir(self) -> None:
        doc = DocIR.from_mapping(
            {
                "s1.p1.r1": "Hello ",
                "s1.p1.r2": "World",
            },
            source_doc_type="docx",
        )

        result = apply_text_edits(
            ApplyTextEditsRequest(
                document=DocumentInput(doc_ir=doc),
                edits=[
                    TextEdit(
                        target_kind="paragraph",
                        target_id=doc.paragraphs[0].node_id,
                        expected_text="Hello World",
                        new_text="Hello Contract World",
                        reason="Expand wording",
                    )
                ],
            )
        )

        self.assertTrue(result.ok)
        self.assertIsNone(result.output_path)
        self.assertIsNone(result.output_bytes)
        self.assertIsNotNone(result.updated_doc_ir)
        self.assertEqual(result.updated_doc_ir.paragraphs[0].text, "Hello Contract World")

    def test_apply_text_edits_accepts_stable_target_id(self) -> None:
        doc = DocIR.from_mapping(
            {
                "s1.p1.r1": "Hello ",
                "s1.p1.r2": "World",
            },
            source_doc_type="docx",
        )
        target_id = doc.paragraphs[0].node_id

        result = apply_text_edits(
            ApplyTextEditsRequest(
                document=DocumentInput(doc_ir=doc),
                edits=[
                    TextEdit(
                        target_kind="paragraph",
                        target_id=target_id,
                        expected_text="Hello World",
                        new_text="Hello Stable World",
                    )
                ],
            )
        )

        self.assertTrue(result.ok)
        self.assertEqual(result.modified_target_ids, [target_id])
        self.assertEqual(result.updated_doc_ir.paragraphs[0].text, "Hello Stable World")
        self.assertEqual(result.updated_doc_ir.paragraphs[0].node_id, target_id)

    def test_apply_text_edits_dry_run_returns_preview_without_native_output(self) -> None:
        doc = DocIR.from_mapping({"s1.p1.r1": "Hello"}, source_doc_type="docx")

        result = apply_text_edits(
            ApplyTextEditsRequest(
                document=DocumentInput(doc_ir=doc),
                edits=[
                    TextEdit(
                        target_kind="run",
                        target_id=doc.paragraphs[0].runs[0].node_id,
                        expected_text="Hello",
                        new_text="Preview",
                    )
                ],
                dry_run=True,
                return_doc_ir=True,
            )
        )

        self.assertTrue(result.ok)
        self.assertEqual(result.edits_applied, 0)
        self.assertIsNone(result.output_path)
        self.assertIsNone(result.output_bytes)
        self.assertEqual(result.updated_doc_ir.paragraphs[0].text, "Preview")
        self.assertEqual(doc.paragraphs[0].text, "Hello")

    def test_validate_text_edits_rejects_missing_target_id(self) -> None:
        doc = DocIR.from_mapping(
            {
                "s1.p1.r1": "Hello",
                "s1.p2.r1": "Other",
            },
            source_doc_type="docx",
        )

        result = apply_text_edits(
            ApplyTextEditsRequest(
                document=DocumentInput(doc_ir=doc),
                edits=[
                    TextEdit(
                        target_kind="paragraph",
                        target_id="missing",
                        expected_text="Hello",
                        new_text="Changed",
                    )
                ],
            )
        )

        self.assertFalse(result.ok)
        self.assertEqual(result.validation.issues[0].code, "target_not_found")

    def test_list_editable_targets_includes_cell_targets(self) -> None:
        doc = DocIR.from_mapping(
            {
                "s1.p1.r1.tbl1.tr1.tc1.p1.r1": "Left",
                "s1.p1.r1.tbl1.tr1.tc2.p1.r1": "Right",
            },
            source_doc_type="docx",
        )

        result = list_editable_targets(
            ListEditableTargetsRequest(
                document=DocumentInput(doc_ir=doc),
                target_kinds=["cell"],
            )
        )

        self.assertEqual(
            [(target.target_kind, target.native_anchor.debug_path, target.current_text) for target in result.targets],
            [
                ("cell", "s1.p1.r1.tbl1.tr1.tc1", "Left"),
                ("cell", "s1.p1.r1.tbl1.tr1.tc2", "Right"),
            ],
        )

    def test_list_editable_targets_does_not_return_all_targets_for_missing_filter(self) -> None:
        doc = DocIR.from_mapping({"s1.p1.r1": "Hello"})

        result = list_editable_targets(
            ListEditableTargetsRequest(
                document=DocumentInput(doc_ir=doc),
                target_ids=["missing"],
            )
        )

        self.assertEqual(result.targets, [])
        self.assertEqual(result.missing_target_ids, ["missing"])

    def test_apply_text_edits_can_replace_doc_ir_cell_text(self) -> None:
        doc = DocIR.from_mapping(
            {
                "s1.p1.r1.tbl1.tr1.tc1.p1.r1": "Left",
                "s1.p1.r1.tbl1.tr1.tc2.p1.r1": "Right",
            },
            source_doc_type="docx",
        )

        result = apply_text_edits(
            ApplyTextEditsRequest(
                document=DocumentInput(doc_ir=doc),
                edits=[
                    TextEdit(
                        target_kind="cell",
                        target_id=doc.paragraphs[0].tables[0].cells[0].node_id,
                        expected_text="Left",
                        new_text="Changed",
                    )
                ],
            )
        )

        self.assertTrue(result.ok)
        cell = doc.paragraphs[0].tables[0].cells[0]
        run = cell.paragraphs[0].runs[0]
        self.assertEqual(result.modified_target_ids, [cell.node_id])
        self.assertEqual(result.modified_run_ids, [run.node_id])
        table = result.updated_doc_ir.paragraphs[0].tables[0]
        self.assertEqual(table.cells[0].text, "Changed")
        self.assertEqual(result.updated_doc_ir.paragraphs[0].text, "Changed\nRight")

    def test_apply_text_edits_replaces_docx_cell_text(self) -> None:
        source_bytes = self._build_sample_table_docx_bytes()
        doc = DocIR.from_file(source_bytes, doc_type="docx")
        result = apply_text_edits(
            ApplyTextEditsRequest(
                document=DocumentInput(
                    source_bytes=source_bytes,
                    source_name="table.docx",
                ),
                edits=[
                    TextEdit(
                        target_kind="cell",
                        target_id=doc.paragraphs[0].tables[0].cells[0].node_id,
                        expected_text="Left",
                        new_text="Changed",
                    )
                ],
                return_doc_ir=True,
            )
        )

        self.assertTrue(result.ok)
        self.assertEqual(result.output_filename, "table_edited.docx")
        self.assertEqual(DocIR.from_file(result.output_bytes).paragraphs[0].tables[0].cells[0].text, "Changed")
        self.assertEqual(result.updated_doc_ir.paragraphs[0].tables[0].cells[0].text, "Changed")

    def test_apply_text_edits_replaces_hwpx_cell_text(self) -> None:
        source_bytes = self._build_sample_table_hwpx_bytes()
        doc = DocIR.from_file(source_bytes, doc_type="hwpx")
        result = apply_text_edits(
            ApplyTextEditsRequest(
                document=DocumentInput(
                    source_bytes=source_bytes,
                    source_name="table.hwpx",
                ),
                edits=[
                    TextEdit(
                        target_kind="cell",
                        target_id=doc.paragraphs[0].tables[0].cells[0].node_id,
                        expected_text="Left",
                        new_text="Changed",
                    )
                ],
                return_doc_ir=True,
            )
        )

        self.assertTrue(result.ok)
        self.assertEqual(result.output_filename, "table_edited.hwpx")
        self.assertEqual(DocIR.from_file(result.output_bytes, doc_type="hwpx").paragraphs[0].tables[0].cells[0].text, "Changed")
        self.assertEqual(result.updated_doc_ir.paragraphs[0].tables[0].cells[0].text, "Changed")

    def test_apply_text_edits_normalizes_hwpx_output_suffix_for_path_backed_writeback(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.hwpx"
            requested_output = Path(tmp_dir) / "sample_edited.docx"
            source.write_bytes(self._build_sample_hwpx_bytes())
            doc = DocIR.from_file(source, doc_type="hwpx")

            result = apply_text_edits(
                ApplyTextEditsRequest(
                    document=DocumentInput(source_path=str(source)),
                    edits=[
                        TextEdit(
                            target_kind="run",
                            target_id=doc.paragraphs[0].runs[1].node_id,
                            expected_text="World",
                            new_text="HWPX",
                            reason="Rename token",
                        )
                    ],
                    output_path=str(requested_output),
                    return_doc_ir=True,
                )
            )

            self.assertTrue(result.ok)
            self.assertEqual(Path(result.output_path).suffix, ".hwpx")
            self.assertEqual(Path(result.output_path).name, "sample_edited.hwpx")
            self.assertTrue(any("adjusted output path" in warning for warning in result.warnings))
            self.assertEqual(result.updated_doc_ir.paragraphs[0].text, "Hello HWPX")

    def test_apply_text_edits_preserves_hwpx_namespace_prefixes_and_declaration(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.hwpx"
            output = Path(tmp_dir) / "sample_edited.hwpx"
            source.write_bytes(self._build_namespaced_hwpx_bytes())
            doc = DocIR.from_file(source, doc_type="hwpx")

            result = apply_text_edits(
                ApplyTextEditsRequest(
                    document=DocumentInput(source_path=str(source)),
                    edits=[
                        TextEdit(
                            target_kind="run",
                            target_id=doc.paragraphs[0].runs[1].node_id,
                            expected_text="World",
                            new_text="HWPX",
                            reason="Rename token",
                        )
                    ],
                    output_path=str(output),
                )
            )

            self.assertTrue(result.ok)
            with zipfile.ZipFile(output) as archive:
                section_xml = archive.read("Contents/section0.xml")

            self.assertTrue(section_xml.startswith(b'<?xml version="1.0" encoding="UTF-8" standalone="yes" ?>'))
            self.assertIn(b'xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core"', section_xml)
            self.assertIn(b"<hc:pt0", section_xml)
            self.assertNotIn(b"xmlns:ns", section_xml)
            self.assertIn(b"HWPX", section_xml)

    def test_render_review_html_accepts_doc_ir_input(self) -> None:
        doc = DocIR.from_mapping({"s1.p1.r1": "Hello"})

        result = render_review_html(
            RenderReviewHtmlRequest(
                document=DocumentInput(doc_ir=doc),
                annotations=[
                    TextAnnotation(
                        target_kind="run",
                        target_id=doc.paragraphs[0].runs[0].node_id,
                        selected_text="Hello",
                        label="Greeting",
                    )
                ],
            )
        )

        self.assertTrue(result.ok)
        self.assertIn("<mark", result.html or "")
        self.assertEqual(result.resolved_annotations[0].selected_text, "Hello")

    def test_render_review_html_accepts_stable_target_id(self) -> None:
        doc = DocIR.from_mapping({"s1.p1.r1": "Hello"})
        run_id = doc.paragraphs[0].runs[0].node_id

        result = render_review_html(
            RenderReviewHtmlRequest(
                document=DocumentInput(doc_ir=doc),
                annotations=[
                    TextAnnotation(
                        target_kind="run",
                        target_id=run_id,
                        selected_text="Hello",
                        label="Greeting",
                    )
                ],
            )
        )

        self.assertTrue(result.ok)
        self.assertEqual(result.resolved_annotations[0].target_id, run_id)

    def test_render_review_html_rejects_ambiguous_selected_text(self) -> None:
        doc = DocIR.from_mapping({"s1.p1.r1": "Hello Hello"})

        result = render_review_html(
            RenderReviewHtmlRequest(
                document=DocumentInput(doc_ir=doc),
                annotations=[
                    TextAnnotation(
                        target_kind="run",
                        target_id=doc.paragraphs[0].runs[0].node_id,
                        selected_text="Hello",
                        label="Ambiguous",
                    )
                ],
            )
        )

        self.assertFalse(result.ok)
        self.assertEqual(result.validation.issues[0].code, "selected_text_ambiguous")


if __name__ == "__main__":
    unittest.main()
