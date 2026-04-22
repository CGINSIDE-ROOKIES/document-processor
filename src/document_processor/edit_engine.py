from __future__ import annotations

from dataclasses import dataclass
import difflib
from io import BytesIO
from pathlib import Path
import re
from typing import BinaryIO, Callable
import tempfile
from xml.etree import ElementTree as ET
import zipfile

from pydantic import BaseModel, Field

from .core import convert_hwp_to_hwpx_bytes
from .io_utils import SourceDocType, TemporarySourcePath, coerce_source_to_supported_value, infer_doc_type
from .models import DocIR, ParagraphIR, RunIR, TableCellIR, TableIR, _anchored_node_id


class EditValidationError(ValueError):
    pass


class RunTextEdit(BaseModel):
    run_id: str
    old_text: str
    new_text: str
    reason: str = ""


class ParagraphTextEdit(BaseModel):
    paragraph_id: str
    old_text: str
    new_text: str
    reason: str = ""


class CellTextEdit(BaseModel):
    cell_id: str
    old_text: str
    new_text: str
    reason: str = ""


EditCommand = RunTextEdit | ParagraphTextEdit | CellTextEdit


class ApplyEditsResult(BaseModel):
    model_config = {"arbitrary_types_allowed": True}

    source_doc_type: str | None = None
    output_path: str | None = None
    output_filename: str | None = None
    output_bytes: bytes | None = None
    updated_doc_ir: DocIR | None = None
    edits_applied: int = 0
    modified_target_ids: list[str] = Field(default_factory=list)
    modified_run_ids: list[str] = Field(default_factory=list)
    warnings: list[str] = Field(default_factory=list)


class _EditableRunRef:
    def __init__(
        self,
        *,
        node_id: str,
        get_text: Callable[[], str],
        set_text: Callable[[str], None],
    ) -> None:
        self.node_id = node_id
        self._get_text = get_text
        self._set_text = set_text

    @property
    def text(self) -> str:
        return self._get_text()

    @text.setter
    def text(self, value: str) -> None:
        self._set_text(value)


class _EditableParagraphRef:
    def __init__(
        self,
        *,
        node_id: str,
        runs: list[_EditableRunRef],
        has_non_run_content: bool = False,
        recompute: Callable[[], None] | None = None,
    ) -> None:
        self.node_id = node_id
        self.runs = runs
        self.has_non_run_content = has_non_run_content
        self._recompute = recompute

    @property
    def text(self) -> str:
        return "".join(run.text for run in self.runs)

    def recompute(self) -> None:
        if self._recompute is not None:
            self._recompute()


class _EditableCellRef:
    def __init__(
        self,
        *,
        node_id: str,
        paragraphs: list[_EditableParagraphRef],
        recompute: Callable[[], None] | None = None,
    ) -> None:
        self.node_id = node_id
        self.paragraphs = paragraphs
        self._recompute = recompute

    @property
    def text(self) -> str:
        return "\n".join(paragraph.text for paragraph in self.paragraphs)

    def recompute(self) -> None:
        if self._recompute is not None:
            self._recompute()


class _EditableDocIndex:
    def __init__(
        self,
        *,
        paragraphs: dict[str, _EditableParagraphRef],
        runs: dict[str, _EditableRunRef],
        cells: dict[str, _EditableCellRef],
        run_to_paragraph: dict[str, _EditableParagraphRef],
    ) -> None:
        self.paragraphs = paragraphs
        self.runs = runs
        self.cells = cells
        self.run_to_paragraph = run_to_paragraph


class _RunSpan(BaseModel):
    start: int
    end: int
    full_start: int
    full_end: int
    run: _EditableRunRef

    model_config = {"arbitrary_types_allowed": True}


def _iter_doc_ir_paragraphs(paragraphs: list[ParagraphIR]):
    for paragraph in paragraphs:
        yield paragraph
        for table in paragraph.tables:
            yield from _iter_doc_ir_table_paragraphs(table)


def _iter_doc_ir_table_paragraphs(table: TableIR):
    for cell in table.cells:
        for paragraph in cell.paragraphs:
            yield paragraph
            for nested_table in paragraph.tables:
                yield from _iter_doc_ir_table_paragraphs(nested_table)


def _build_doc_ir_index(doc: DocIR) -> _EditableDocIndex:
    doc.ensure_node_identity()
    paragraphs: dict[str, _EditableParagraphRef] = {}
    runs: dict[str, _EditableRunRef] = {}
    cells: dict[str, _EditableCellRef] = {}
    run_to_paragraph: dict[str, _EditableParagraphRef] = {}

    def register_paragraph(
        paragraph: ParagraphIR,
        *,
        recompute_after: Callable[[], None] | None = None,
    ) -> _EditableParagraphRef:
        run_refs: list[_EditableRunRef] = []

        def recompute() -> None:
            paragraph.recompute_text()
            if recompute_after is not None:
                recompute_after()

        paragraph_ref = _EditableParagraphRef(
            node_id=paragraph.node_id,
            runs=run_refs,
            has_non_run_content=bool(paragraph.images or paragraph.tables),
            recompute=recompute,
        )
        paragraphs[paragraph.node_id] = paragraph_ref
        for run in paragraph.runs:
            run_ref = _EditableRunRef(
                node_id=run.node_id,
                get_text=lambda node=run: node.text,
                set_text=lambda value, node=run: setattr(node, "text", value),
            )
            run_refs.append(run_ref)
            runs[run.node_id] = run_ref
            run_to_paragraph[run.node_id] = paragraph_ref
        return paragraph_ref

    def walk_table(table: TableIR, *, recompute_after: Callable[[], None] | None) -> None:
        for cell in table.cells:
            cell_paragraph_refs: list[_EditableParagraphRef] = []

            def recompute_cell(node: TableCellIR = cell) -> None:
                node.recompute_text()
                if recompute_after is not None:
                    recompute_after()

            cell_ref = _EditableCellRef(
                node_id=cell.node_id,
                paragraphs=cell_paragraph_refs,
                recompute=recompute_cell,
            )
            cells[cell.node_id] = cell_ref

            for cell_paragraph in cell.paragraphs:
                paragraph_ref = register_paragraph(
                    cell_paragraph,
                    recompute_after=cell_ref.recompute,
                )
                cell_paragraph_refs.append(paragraph_ref)
                for nested_table in cell_paragraph.tables:
                    walk_table(nested_table, recompute_after=paragraph_ref.recompute)

    for paragraph in doc.paragraphs:
        paragraph_ref = register_paragraph(paragraph)
        for table in paragraph.tables:
            walk_table(table, recompute_after=paragraph_ref.recompute)

    return _EditableDocIndex(paragraphs=paragraphs, runs=runs, cells=cells, run_to_paragraph=run_to_paragraph)


def _iter_docx_blocks(doc):
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    iter_inner_content = getattr(doc, "iter_inner_content", None)
    if callable(iter_inner_content):
        yield from iter_inner_content()
        return

    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def _iter_docx_blocks_from_element(parent, element):
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _build_docx_index(doc) -> _EditableDocIndex:
    paragraphs: dict[str, _EditableParagraphRef] = {}
    runs: dict[str, _EditableRunRef] = {}
    cells: dict[str, _EditableCellRef] = {}
    run_to_paragraph: dict[str, _EditableParagraphRef] = {}

    def register_paragraph(paragraph, paragraph_path: str, *, has_non_run_content: bool = False) -> _EditableParagraphRef:
        run_refs: list[_EditableRunRef] = []
        paragraph_node_id = _anchored_node_id("paragraph", paragraph_path)
        paragraph_ref = _EditableParagraphRef(
            node_id=paragraph_node_id,
            runs=run_refs,
            has_non_run_content=has_non_run_content,
        )
        paragraphs[paragraph_node_id] = paragraph_ref
        for run_index, run in enumerate(paragraph.runs, start=1):
            run_path = f"{paragraph_path}.r{run_index}"
            run_node_id = _anchored_node_id("run", run_path)
            run_ref = _EditableRunRef(
                node_id=run_node_id,
                get_text=lambda node=run: node.text or "",
                set_text=lambda value, node=run: setattr(node, "text", value),
            )
            run_refs.append(run_ref)
            runs[run_node_id] = run_ref
            run_to_paragraph[run_node_id] = paragraph_ref
        return paragraph_ref

    def walk_table(table, table_base: str) -> None:
        for tr_idx, row in enumerate(table.rows, start=1):
            for tc_idx, cell in enumerate(row.cells, start=1):
                cell_path = f"{table_base}.tr{tr_idx}.tc{tc_idx}"
                cell_node_id = _anchored_node_id("cell", cell_path)
                cell_paragraph_refs: list[_EditableParagraphRef] = []
                cells[cell_node_id] = _EditableCellRef(node_id=cell_node_id, paragraphs=cell_paragraph_refs)
                cp_idx = 0
                current_paragraph_path: str | None = None
                nested_table_counter_by_paragraph: dict[str, int] = {}

                for block in _iter_docx_blocks_from_element(cell, cell._tc):
                    if block.__class__.__name__ == "Paragraph":
                        cp_idx += 1
                        current_paragraph_path = f"{cell_path}.p{cp_idx}"
                        cell_paragraph_refs.append(register_paragraph(block, current_paragraph_path))
                        continue

                    if block.__class__.__name__ != "Table":
                        continue

                    if current_paragraph_path is None:
                        cp_idx += 1
                        current_paragraph_path = f"{cell_path}.p{cp_idx}"
                        paragraph_node_id = _anchored_node_id("paragraph", current_paragraph_path)
                        paragraph_ref = _EditableParagraphRef(
                            node_id=paragraph_node_id,
                            runs=[],
                            has_non_run_content=True,
                        )
                        paragraphs[paragraph_node_id] = paragraph_ref
                        cell_paragraph_refs.append(paragraph_ref)
                    else:
                        paragraphs[_anchored_node_id("paragraph", current_paragraph_path)].has_non_run_content = True

                    tbl_counter = nested_table_counter_by_paragraph.get(current_paragraph_path, 0) + 1
                    nested_table_counter_by_paragraph[current_paragraph_path] = tbl_counter
                    nested_table_base = f"{current_paragraph_path}.tbl{tbl_counter}"
                    walk_table(block, nested_table_base)

    p_idx = 0
    tbl_counter = 0
    for block in _iter_docx_blocks(doc):
        if block.__class__.__name__ == "Paragraph":
            p_idx += 1
            register_paragraph(block, f"s1.p{p_idx}")
            continue

        if block.__class__.__name__ != "Table":
            continue

        tbl_counter += 1
        p_idx += 1
        walk_table(block, f"s1.p{p_idx}.r1.tbl{tbl_counter}")

    return _EditableDocIndex(paragraphs=paragraphs, runs=runs, cells=cells, run_to_paragraph=run_to_paragraph)


_SECTION_NAME_RE = re.compile(r"^Contents/section(\d+)\.xml$")
_HP_NS = "http://www.hancom.co.kr/hwpml/2011/paragraph"
_HP = f"{{{_HP_NS}}}"
_XML_PREFIX_AND_ROOT_RE = re.compile(
    rb"^(?P<prefix>\s*(?:<\?xml[^>]*\?>\s*)?)(?P<root_open><[^!?][^>]*>)",
    re.DOTALL,
)


@dataclass
class _EditableHwpxSection:
    name: str
    root: ET.Element
    xml_prefix: bytes
    original_root_open: bytes
    namespaces: list[tuple[str, str]]


def _run_text(run_el: ET.Element) -> str:
    return "".join("".join(node.itertext()) for node in run_el.findall(f"{_HP}t"))


def _set_hwpx_run_text(run_el: ET.Element, new_text: str) -> None:
    for node in list(run_el):
        if node.tag == f"{_HP}t":
            run_el.remove(node)
    text_el = ET.SubElement(run_el, f"{_HP}t")
    text_el.text = new_text


def _iter_section_paragraphs(section_root: ET.Element) -> list[ET.Element]:
    return section_root.findall(f"{_HP}p")


def _iter_paragraph_tables(paragraph_el: ET.Element) -> list[ET.Element]:
    return paragraph_el.findall(f"{_HP}run/{_HP}tbl")


def _iter_cell_paragraphs(cell_el: ET.Element) -> list[ET.Element]:
    direct = cell_el.findall(f"{_HP}subList/{_HP}p")
    if direct:
        return direct
    return cell_el.findall(f".//{_HP}p")


def _safe_int(value: str | None) -> int | None:
    if value is None:
        return None
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _logical_table_cells(row_el: ET.Element) -> list[tuple[int, ET.Element]]:
    logical_cells: list[tuple[int, ET.Element]] = []
    fallback_col = 1
    for cell_el in row_el.findall(f"{_HP}tc"):
        cell_addr = cell_el.find(f"{_HP}cellAddr")
        col_addr = _safe_int(cell_addr.get("colAddr")) if cell_addr is not None else None
        logical_col = (col_addr + 1) if col_addr is not None else fallback_col
        logical_cells.append((logical_col, cell_el))

        cell_span = cell_el.find(f"{_HP}cellSpan")
        colspan = _safe_int(cell_span.get("colSpan")) if cell_span is not None else None
        fallback_col = max(fallback_col, logical_col + max(colspan or 1, 1))
    return logical_cells


def _collect_xml_namespaces(xml_bytes: bytes) -> list[tuple[str, str]]:
    namespaces: list[tuple[str, str]] = []
    for _event, item in ET.iterparse(BytesIO(xml_bytes), events=("start-ns",)):
        if item not in namespaces:
            namespaces.append(item)
    return namespaces


def _split_xml_prefix_and_root_open(xml_bytes: bytes) -> tuple[bytes, bytes] | None:
    match = _XML_PREFIX_AND_ROOT_RE.match(xml_bytes)
    if match is None:
        return None
    return match.group("prefix"), match.group("root_open")


def _serialize_hwpx_section(section: _EditableHwpxSection) -> bytes:
    for prefix, uri in section.namespaces:
        ET.register_namespace(prefix, uri)

    serialized = ET.tostring(section.root, encoding="utf-8", xml_declaration=False)
    generated_parts = _split_xml_prefix_and_root_open(serialized)
    if generated_parts is None:
        return serialized if not section.xml_prefix else section.xml_prefix + serialized

    _generated_prefix, generated_root_open = generated_parts
    generated_body = serialized[len(generated_root_open) :]
    return section.xml_prefix + section.original_root_open + generated_body


class _EditableHwpxArchive:
    def __init__(
        self,
        *,
        source_path: Path,
        source_bytes: bytes,
        section_entries: list[_EditableHwpxSection],
    ) -> None:
        self.source_path = source_path
        self.source_bytes = source_bytes
        self.section_entries = section_entries

    @staticmethod
    def _load_section_entries(source_bytes: bytes) -> list[_EditableHwpxSection]:
        with zipfile.ZipFile(BytesIO(source_bytes)) as archive:
            return sorted(
                [
                    _EditableHwpxSection(
                        name=name,
                        root=ET.fromstring(section_bytes := archive.read(name)),
                        xml_prefix=(split_parts[0] if (split_parts := _split_xml_prefix_and_root_open(section_bytes)) else b""),
                        original_root_open=(split_parts[1] if split_parts else b""),
                        namespaces=_collect_xml_namespaces(section_bytes),
                    )
                    for name in archive.namelist()
                    if _SECTION_NAME_RE.match(name)
                ],
                key=lambda item: int(_SECTION_NAME_RE.match(item.name).group(1)),
            )

    @classmethod
    def open(cls, source_path: str | Path) -> "_EditableHwpxArchive":
        path = Path(source_path)
        source_bytes = path.read_bytes()
        section_entries = cls._load_section_entries(source_bytes)
        return cls(source_path=path, source_bytes=source_bytes, section_entries=section_entries)

    @classmethod
    def from_bytes(
        cls,
        source_bytes: bytes,
        *,
        source_path: str | Path | None = None,
    ) -> "_EditableHwpxArchive":
        path = Path(source_path) if source_path is not None else Path("converted.hwpx")
        section_entries = cls._load_section_entries(source_bytes)
        return cls(source_path=path, source_bytes=source_bytes, section_entries=section_entries)

    def write_to(self, output_path: str | Path) -> None:
        section_bytes = {
            section.name: _serialize_hwpx_section(section)
            for section in self.section_entries
        }

        output = Path(output_path)
        with zipfile.ZipFile(BytesIO(self.source_bytes), "r") as source_archive:
            with zipfile.ZipFile(output, "w") as target_archive:
                for info in source_archive.infolist():
                    data = section_bytes.get(info.filename, source_archive.read(info.filename))
                    target_archive.writestr(info, data)


def _build_hwpx_index(archive: _EditableHwpxArchive) -> _EditableDocIndex:
    paragraphs: dict[str, _EditableParagraphRef] = {}
    runs: dict[str, _EditableRunRef] = {}
    cells: dict[str, _EditableCellRef] = {}
    run_to_paragraph: dict[str, _EditableParagraphRef] = {}

    def register_paragraph(paragraph_el: ET.Element, paragraph_path: str) -> _EditableParagraphRef:
        run_elements = paragraph_el.findall(f"{_HP}run")
        paragraph_node_id = _anchored_node_id("paragraph", paragraph_path)
        paragraph_ref = _EditableParagraphRef(
            node_id=paragraph_node_id,
            runs=[],
            has_non_run_content=bool(_iter_paragraph_tables(paragraph_el)),
        )
        paragraphs[paragraph_node_id] = paragraph_ref
        if not run_elements:
            return paragraph_ref

        for run_index, run_el in enumerate(run_elements, start=1):
            run_path = f"{paragraph_path}.r{run_index}"
            run_node_id = _anchored_node_id("run", run_path)
            run_ref = _EditableRunRef(
                node_id=run_node_id,
                get_text=lambda node=run_el: _run_text(node),
                set_text=lambda value, node=run_el: _set_hwpx_run_text(node, value),
            )
            paragraph_ref.runs.append(run_ref)
            runs[run_node_id] = run_ref
            run_to_paragraph[run_node_id] = paragraph_ref
        return paragraph_ref

    def walk_table(table_el: ET.Element, table_base: str) -> None:
        for tr_idx, row_el in enumerate(table_el.findall(f"{_HP}tr"), start=1):
            for tc_idx, cell_el in _logical_table_cells(row_el):
                cell_path = f"{table_base}.tr{tr_idx}.tc{tc_idx}"
                cell_node_id = _anchored_node_id("cell", cell_path)
                cell_paragraph_refs: list[_EditableParagraphRef] = []
                cells[cell_node_id] = _EditableCellRef(node_id=cell_node_id, paragraphs=cell_paragraph_refs)
                cell_paragraphs = _iter_cell_paragraphs(cell_el)
                if not cell_paragraphs:
                    paragraph_path = f"{cell_path}.p1"
                    paragraph_node_id = _anchored_node_id("paragraph", paragraph_path)
                    paragraph_ref = _EditableParagraphRef(
                        node_id=paragraph_node_id,
                        runs=[],
                        has_non_run_content=False,
                    )
                    paragraphs[paragraph_node_id] = paragraph_ref
                    cell_paragraph_refs.append(paragraph_ref)
                    continue

                for cp_idx, paragraph_el in enumerate(cell_paragraphs, start=1):
                    paragraph_path = f"{cell_path}.p{cp_idx}"
                    paragraph_ref = register_paragraph(paragraph_el, paragraph_path)
                    cell_paragraph_refs.append(paragraph_ref)
                    for nested_index, nested_table in enumerate(_iter_paragraph_tables(paragraph_el), start=1):
                        paragraph_ref.has_non_run_content = True
                        walk_table(nested_table, f"{paragraph_path}.tbl{nested_index}")

    for section_index, section in enumerate(archive.section_entries, start=1):
        for paragraph_index, paragraph_el in enumerate(_iter_section_paragraphs(section.root), start=1):
            paragraph_id = f"s{section_index}.p{paragraph_index}"
            paragraph_ref = register_paragraph(paragraph_el, paragraph_id)
            for table_index, table_el in enumerate(_iter_paragraph_tables(paragraph_el), start=1):
                paragraph_ref.has_non_run_content = True
                walk_table(table_el, f"{paragraph_id}.r1.tbl{table_index}")

    return _EditableDocIndex(paragraphs=paragraphs, runs=runs, cells=cells, run_to_paragraph=run_to_paragraph)


def _build_run_spans(paragraph: _EditableParagraphRef) -> list[_RunSpan]:
    spans: list[_RunSpan] = []
    cursor = 0
    for run in paragraph.runs:
        length = len(run.text)
        spans.append(
            _RunSpan(
                start=cursor,
                end=cursor + length,
                full_start=cursor,
                full_end=cursor + length,
                run=run,
            )
        )
        cursor += length
    return spans


def _clip_run_spans(spans: list[_RunSpan], i1: int, i2: int) -> list[_RunSpan]:
    is_insert = i1 == i2
    clipped: list[_RunSpan] = []
    for span in spans:
        if is_insert:
            if span.end < i1 or span.start > i2:
                continue
        else:
            if span.end <= i1 or span.start >= i2:
                continue
        clipped.append(
            _RunSpan(
                start=span.start if is_insert else max(span.start, i1),
                end=span.end if is_insert else min(span.end, i2),
                full_start=span.full_start,
                full_end=span.full_end,
                run=span.run,
            )
        )
    if is_insert and len(clipped) > 1:
        preceding = [span for span in clipped if span.end == i1]
        if preceding:
            clipped = preceding[:1]
    return clipped


def _append_unique(values: list[str], value: str) -> None:
    if value not in values:
        values.append(value)


def _apply_to_run(run: _EditableRunRef, new_text: str, local_start: int, local_end: int) -> None:
    current = run.text
    run.text = current[:local_start] + new_text + current[local_end:]


def _apply_to_run_with_offsets(
    run: _EditableRunRef,
    new_text: str,
    local_start: int,
    local_end: int,
    offset_deltas: dict[str, int],
) -> None:
    delta = offset_deltas.get(run.node_id, 0)
    actual_start = local_start + delta
    actual_end = local_end + delta
    _apply_to_run(run, new_text, actual_start, actual_end)
    offset_deltas[run.node_id] = delta + len(new_text) - (local_end - local_start)


def _apply_multi_run(
    orig_sub: str,
    new_sub: str,
    spans: list[_RunSpan],
    result: ApplyEditsResult,
    offset_deltas: dict[str, int],
    *,
    base_offset: int,
    depth: int = 0,
) -> None:
    if depth == 0:
        matcher = difflib.SequenceMatcher(None, orig_sub, new_sub, autojunk=False)
        for tag, a1, a2, b1, b2 in matcher.get_opcodes():
            if tag == "equal":
                continue
            abs_a1 = base_offset + a1
            abs_a2 = base_offset + a2
            sub_spans = _clip_run_spans(spans, abs_a1, abs_a2)
            if not sub_spans:
                continue
            if len(sub_spans) == 1:
                span = sub_spans[0]
                local_start = abs_a1 - span.full_start
                local_end = abs_a2 - span.full_start
                _apply_to_run_with_offsets(span.run, new_sub[b1:b2], local_start, local_end, offset_deltas)
                _append_unique(result.modified_run_ids, span.run.node_id)
                continue
            _apply_multi_run(
                orig_sub[a1:a2],
                new_sub[b1:b2],
                sub_spans,
                result,
                offset_deltas,
                base_offset=abs_a1,
                depth=1,
            )
        return

    first = spans[0]
    _apply_to_run_with_offsets(
        first.run,
        new_sub,
        first.start - first.full_start,
        first.end - first.full_start,
        offset_deltas,
    )
    _append_unique(result.modified_run_ids, first.run.node_id)
    for span in spans[1:]:
        _apply_to_run_with_offsets(
            span.run,
            "",
            span.start - span.full_start,
            span.end - span.full_start,
            offset_deltas,
        )
        _append_unique(result.modified_run_ids, span.run.node_id)
    result.warnings.append(
        "Multi-run fallback used for "
        f"{[span.run.node_id for span in spans]}: all replacement text assigned to {first.run.node_id}"
    )


def _validate_paragraph_edit(index: _EditableDocIndex, edit: ParagraphTextEdit) -> _EditableParagraphRef:
    paragraph = index.paragraphs.get(edit.paragraph_id)
    if paragraph is None:
        raise EditValidationError(f"Paragraph does not exist: {edit.paragraph_id}")
    if paragraph.has_non_run_content:
        raise EditValidationError(
            f"Paragraph edit targets unsupported mixed content (tables/images): {edit.paragraph_id}"
        )
    if paragraph.text != edit.old_text:
        raise EditValidationError(
            f"Paragraph text mismatch for {edit.paragraph_id}: expected {edit.old_text!r}, got {paragraph.text!r}"
        )
    return paragraph


def _validate_run_edit(index: _EditableDocIndex, edit: RunTextEdit) -> _EditableRunRef:
    run = index.runs.get(edit.run_id)
    if run is None:
        raise EditValidationError(f"Run does not exist: {edit.run_id}")
    if run.text != edit.old_text:
        raise EditValidationError(
            f"Run text mismatch for {edit.run_id}: expected {edit.old_text!r}, got {run.text!r}"
        )
    return run


def _validate_cell_edit(index: _EditableDocIndex, edit: CellTextEdit) -> _EditableCellRef:
    cell = index.cells.get(edit.cell_id)
    if cell is None:
        raise EditValidationError(f"Cell does not exist: {edit.cell_id}")
    if any(paragraph.has_non_run_content for paragraph in cell.paragraphs):
        raise EditValidationError(
            f"Cell edit targets unsupported mixed content (nested tables/images): {edit.cell_id}"
        )
    if not cell.paragraphs or any(not paragraph.runs for paragraph in cell.paragraphs):
        raise EditValidationError(f"Cell does not contain editable text runs: {edit.cell_id}")
    if cell.text != edit.old_text:
        raise EditValidationError(
            f"Cell text mismatch for {edit.cell_id}: expected {edit.old_text!r}, got {cell.text!r}"
        )
    expected_paragraphs = len(cell.paragraphs)
    new_paragraphs = len(edit.new_text.split("\n"))
    if new_paragraphs != expected_paragraphs:
        raise EditValidationError(
            f"Cell text replacement for {edit.cell_id} must preserve paragraph count: "
            f"expected {expected_paragraphs} line(s), got {new_paragraphs}."
        )
    return cell


def _replace_paragraph_text(
    paragraph: _EditableParagraphRef,
    new_text: str,
    result: ApplyEditsResult,
) -> None:
    spans = _build_run_spans(paragraph)
    original = paragraph.text
    if len(spans) == 1:
        run = spans[0].run
        run.text = new_text
        _append_unique(result.modified_run_ids, run.node_id)
        paragraph.recompute()
        return

    offset_deltas: dict[str, int] = {}
    matcher = difflib.SequenceMatcher(None, original, new_text, autojunk=False)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        affected = _clip_run_spans(spans, i1, i2)
        if not affected:
            continue
        if len(affected) == 1:
            span = affected[0]
            local_start = i1 - span.full_start
            local_end = i2 - span.full_start
            _apply_to_run_with_offsets(span.run, new_text[j1:j2], local_start, local_end, offset_deltas)
            _append_unique(result.modified_run_ids, span.run.node_id)
            continue
        _apply_multi_run(
            original[i1:i2],
            new_text[j1:j2],
            affected,
            result,
            offset_deltas,
            base_offset=i1,
        )
    paragraph.recompute()


def _apply_single_edit(index: _EditableDocIndex, edit: EditCommand, result: ApplyEditsResult) -> None:
    if isinstance(edit, RunTextEdit):
        run = _validate_run_edit(index, edit)
        run.text = edit.new_text
        paragraph = index.run_to_paragraph.get(edit.run_id)
        if paragraph is not None:
            paragraph.recompute()
        _append_unique(result.modified_target_ids, edit.run_id)
        _append_unique(result.modified_run_ids, edit.run_id)
        result.edits_applied += 1
        return

    if isinstance(edit, ParagraphTextEdit):
        paragraph = _validate_paragraph_edit(index, edit)
        _replace_paragraph_text(paragraph, edit.new_text, result)
        _append_unique(result.modified_target_ids, edit.paragraph_id)
        result.edits_applied += 1
        return

    cell = _validate_cell_edit(index, edit)
    for paragraph, new_paragraph_text in zip(cell.paragraphs, edit.new_text.split("\n"), strict=True):
        _replace_paragraph_text(paragraph, new_paragraph_text, result)
    cell.recompute()
    _append_unique(result.modified_target_ids, edit.cell_id)
    result.edits_applied += 1


def validate_edit_commands(doc: DocIR, edits: list[EditCommand]) -> None:
    index = _build_doc_ir_index(doc)
    for edit in edits:
        if isinstance(edit, RunTextEdit):
            _validate_run_edit(index, edit)
        elif isinstance(edit, ParagraphTextEdit):
            _validate_paragraph_edit(index, edit)
        else:
            _validate_cell_edit(index, edit)


def apply_edits_to_doc_ir(doc: DocIR, edits: list[EditCommand]) -> tuple[DocIR, ApplyEditsResult]:
    updated = doc.model_copy(deep=True)
    index = _build_doc_ir_index(updated)
    result = ApplyEditsResult(source_doc_type=updated.source_doc_type)
    for edit in edits:
        _apply_single_edit(index, edit, result)
    result.updated_doc_ir = updated
    return updated, result


def _default_output_path(source_path: Path, *, output_suffix: str | None = None) -> Path:
    suffix = output_suffix if output_suffix is not None else source_path.suffix
    return source_path.with_name(f"{source_path.stem}_edited{suffix}")


def _expected_writeback_suffix(source_doc_type: str | None) -> str | None:
    if source_doc_type == "docx":
        return ".docx"
    if source_doc_type in {"hwpx", "hwp"}:
        return ".hwpx"
    return None


def _normalize_output_path_for_source_doc_type(
    target_path: Path,
    *,
    source_doc_type: str | None,
    result: ApplyEditsResult,
) -> Path:
    expected_suffix = _expected_writeback_suffix(source_doc_type)
    if expected_suffix is None or target_path.suffix.lower() == expected_suffix:
        return target_path

    adjusted_target_path = target_path.with_suffix(expected_suffix)
    if source_doc_type == "hwp":
        result.warnings.append(
            f"HWP sources are written back as HWPX; adjusted output path to {adjusted_target_path}."
        )
    else:
        result.warnings.append(
            f"{str(source_doc_type).upper()} write-back keeps the native {expected_suffix} format; "
            f"adjusted output path to {adjusted_target_path}."
        )
    return adjusted_target_path


def _same_path(left: Path, right: Path) -> bool:
    return left.expanduser().resolve(strict=False) == right.expanduser().resolve(strict=False)


def _default_output_filename(
    *,
    source_name: str | None,
    source_doc_type: str | None,
) -> str:
    source_doc_type = source_doc_type or "docx"
    if source_name:
        source = Path(source_name)
    else:
        suffix = f".{source_doc_type}" if source_doc_type != "hwp" else ".hwp"
        source = Path(f"document{suffix}")

    suffix = ".hwpx" if source_doc_type == "hwp" else (source.suffix or f".{source_doc_type}")
    stem = source.stem if source.suffix else source.name
    return f"{stem}_edited{suffix}"


def _source_suffix_for_doc_type(doc_type: SourceDocType | str) -> str:
    return {
        "docx": ".docx",
        "hwpx": ".hwpx",
        "hwp": ".hwp",
        "pdf": ".pdf",
    }.get(str(doc_type), ".bin")


def _resolve_bytes_doc_type(
    source_bytes: bytes,
    *,
    doc_type: SourceDocType = "auto",
    source_name: str | None = None,
) -> str:
    if doc_type != "auto":
        return doc_type
    if source_name:
        try:
            return infer_doc_type(Path(source_name), "auto")
        except ValueError:
            pass
    return infer_doc_type(source_bytes, "auto")


def apply_edits_to_bytes(
    source_bytes: bytes,
    edits: list[EditCommand],
    *,
    doc_type: SourceDocType = "auto",
    source_name: str | None = None,
    output_filename: str | None = None,
) -> ApplyEditsResult:
    resolved_doc_type = _resolve_bytes_doc_type(
        source_bytes,
        doc_type=doc_type,
        source_name=source_name,
    )

    with TemporarySourcePath(source_bytes, suffix=_source_suffix_for_doc_type(resolved_doc_type)) as source_path:
        default_filename = _default_output_filename(
            source_name=source_name,
            source_doc_type=resolved_doc_type,
        )
        chosen_filename = output_filename or default_filename
        with tempfile.TemporaryDirectory() as tmp_dir:
            target_path = Path(tmp_dir) / chosen_filename
            result = apply_edits_to_file(source_path, edits, output_path=target_path)
            output_path = Path(result.output_path) if result.output_path is not None else target_path
            result.output_bytes = output_path.read_bytes()
            result.output_filename = output_path.name
            result.output_path = None
            return result


def apply_edits_to_source(
    source: DocIR | str | Path | bytes | BinaryIO,
    edits: list[EditCommand],
    *,
    doc_type: SourceDocType = "auto",
    source_name: str | None = None,
    output_path: str | Path | None = None,
    output_filename: str | None = None,
) -> ApplyEditsResult:
    if isinstance(source, DocIR):
        updated, result = apply_edits_to_doc_ir(source, edits)
        result.updated_doc_ir = updated
        return result

    if output_path is not None and output_filename is not None:
        raise ValueError("Specify either output_path or output_filename, not both.")

    if isinstance(source, (str, Path)):
        resolved_output_path = output_path
        if resolved_output_path is None and output_filename is not None:
            resolved_output_path = Path(source).with_name(output_filename)
        result = apply_edits_to_file(source, edits, output_path=resolved_output_path)
        if result.output_path is not None:
            result.output_filename = Path(result.output_path).name
        return result

    source_bytes = coerce_source_to_supported_value(source, doc_type=infer_doc_type(source, doc_type))
    if not isinstance(source_bytes, bytes):
        raise TypeError("Expected bytes-like source after coercion.")
    return apply_edits_to_bytes(
        source_bytes,
        edits,
        doc_type=doc_type,
        source_name=source_name,
        output_filename=output_filename,
    )


def apply_edits_to_file(
    source_path: str | Path,
    edits: list[EditCommand],
    *,
    output_path: str | Path | None = None,
) -> ApplyEditsResult:
    source = Path(source_path)
    doc = DocIR.from_file(source)
    result = ApplyEditsResult(source_doc_type=doc.source_doc_type)
    target_suffix = ".hwpx" if doc.source_doc_type == "hwp" else None
    target_path = Path(output_path) if output_path is not None else _default_output_path(source, output_suffix=target_suffix)
    target_path = _normalize_output_path_for_source_doc_type(
        target_path,
        source_doc_type=doc.source_doc_type,
        result=result,
    )

    if _same_path(source, target_path):
        raise EditValidationError(
            f"Refusing to overwrite source file {source}; choose a different output path."
        )

    if doc.source_doc_type == "docx":
        from docx import Document as load_docx

        native_doc = load_docx(str(source))
        index = _build_docx_index(native_doc)
        for edit in edits:
            _apply_single_edit(index, edit, result)
        native_doc.save(str(target_path))
    elif doc.source_doc_type == "hwpx":
        archive = _EditableHwpxArchive.open(source)
        index = _build_hwpx_index(archive)
        for edit in edits:
            _apply_single_edit(index, edit, result)
        archive.write_to(target_path)
    elif doc.source_doc_type == "hwp":
        archive = _EditableHwpxArchive.from_bytes(
            convert_hwp_to_hwpx_bytes(source),
            source_path=source.with_suffix(".hwpx"),
        )
        index = _build_hwpx_index(archive)
        for edit in edits:
            _apply_single_edit(index, edit, result)
        archive.write_to(target_path)
    else:
        raise EditValidationError(
            f"Native write-back is currently supported only for docx/hwp/hwpx, got {doc.source_doc_type!r}."
        )

    result.output_path = str(target_path)
    result.output_filename = target_path.name
    return result


__all__ = [
    "ApplyEditsResult",
    "CellTextEdit",
    "EditValidationError",
    "ParagraphTextEdit",
    "RunTextEdit",
    "apply_edits_to_bytes",
    "apply_edits_to_doc_ir",
    "apply_edits_to_file",
    "apply_edits_to_source",
    "validate_edit_commands",
]
