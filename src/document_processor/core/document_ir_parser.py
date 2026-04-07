"""Direct DOCX/HWPX file parsing into structural document IR."""

from __future__ import annotations

import hashlib
import mimetypes
import struct
from io import BytesIO
from pathlib import Path
from typing import TYPE_CHECKING, Any, Callable, Literal
from xml.etree import ElementTree as ET
import zipfile

from ..builder import normalize_text_default
from ..io_utils import coerce_source_to_supported_value
from ..models import DocIR, ImageAsset, ImageIR, ParagraphIR, RunIR, TableCellIR, TableIR
from .docx_structured_exporter import _iter_blocks, _iter_blocks_from_element, _load_docx_source
from .hwp_converter import convert_hwp_to_hwpx_bytes
from .hwpx_structured_exporter import _HP, _logical_table_cells, _paragraph_text, _run_text, _safe_int, _section_roots_from_bytes

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
    from hwpx import HwpxDocument


DocType = Literal["hwp", "hwpx", "docx"]

_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_DOCX_EMBED_ATTR = f"{{{_REL_NS}}}embed"
_HC_IMG_TAG = "img"
_EMU_PER_PT = 12700.0
_HWPUNIT_PER_PT = 100.0


def _emu_to_pt(value: str | int | None) -> float | None:
    if value is None:
        return None
    try:
        return int(value) / _EMU_PER_PT
    except (TypeError, ValueError):
        return None


def _hwpunit_to_pt(value: str | int | None) -> float | None:
    if value is None:
        return None
    try:
        return int(value) / _HWPUNIT_PER_PT
    except (TypeError, ValueError):
        return None


def _image_dimensions_from_bytes(data: bytes) -> tuple[int | None, int | None]:
    if data.startswith(b"\x89PNG\r\n\x1a\n") and len(data) >= 24:
        return struct.unpack(">II", data[16:24])

    if data.startswith((b"GIF87a", b"GIF89a")) and len(data) >= 10:
        return struct.unpack("<HH", data[6:10])

    if len(data) >= 4 and data[:2] == b"\xff\xd8":
        offset = 2
        while offset + 9 < len(data):
            if data[offset] != 0xFF:
                offset += 1
                continue
            marker = data[offset + 1]
            offset += 2
            if marker in (0xD8, 0xD9):
                continue
            if offset + 2 > len(data):
                break
            segment_length = int.from_bytes(data[offset:offset + 2], "big")
            if segment_length < 2 or offset + segment_length > len(data):
                break
            if marker in {
                0xC0, 0xC1, 0xC2, 0xC3,
                0xC5, 0xC6, 0xC7,
                0xC9, 0xCA, 0xCB,
                0xCD, 0xCE, 0xCF,
            } and offset + 7 < len(data):
                height = int.from_bytes(data[offset + 3:offset + 5], "big")
                width = int.from_bytes(data[offset + 5:offset + 7], "big")
                return width, height
            offset += segment_length

    return None, None


def _mime_type_for_filename(filename: str | None) -> str:
    if not filename:
        return "application/octet-stream"
    mime_type, _ = mimetypes.guess_type(filename)
    return mime_type or "application/octet-stream"


def _register_image_asset(
    assets: dict[str, ImageAsset],
    asset_lookup: dict[tuple[str, str], str],
    *,
    data: bytes,
    mime_type: str,
    filename: str | None,
) -> str:
    digest = hashlib.sha1(data).hexdigest()
    cache_key = (digest, mime_type)
    existing = asset_lookup.get(cache_key)
    if existing is not None:
        return existing

    image_id = f"img{len(assets) + 1}"
    intrinsic_width_px, intrinsic_height_px = _image_dimensions_from_bytes(data)
    assets[image_id] = ImageAsset.from_bytes(
        image_id=image_id,
        data=data,
        mime_type=mime_type,
        filename=filename,
        intrinsic_width_px=intrinsic_width_px,
        intrinsic_height_px=intrinsic_height_px,
    )
    asset_lookup[cache_key] = image_id
    return image_id


def _resolve_doc_metadata(
    *,
    source_path: str | Path | None,
    source_doc_type: str,
    metadata: dict[str, Any] | None,
    doc_id: str | None,
) -> tuple[str | None, str | None, dict[str, Any]]:
    resolved_source_path = str(source_path) if source_path is not None else None
    resolved_doc_id = doc_id
    if resolved_doc_id is None and source_path is not None:
        resolved_doc_id = Path(source_path).stem
    return resolved_doc_id, resolved_source_path, metadata or {}


def _parse_docx_run_images(
    run,
    *,
    paragraph_id: str,
    image_counter: int,
    assets: dict[str, ImageAsset],
    asset_lookup: dict[tuple[str, str], str],
) -> tuple[list[ImageIR], int]:
    images: list[ImageIR] = []

    for element in run._r.iter():
        tag = getattr(element, "tag", "")
        if not isinstance(tag, str) or not tag.endswith("}blip"):
            continue

        rel_id = element.get(_DOCX_EMBED_ATTR)
        if not rel_id:
            continue

        image_part = run.part.related_parts.get(rel_id)
        if image_part is None or not hasattr(image_part, "blob"):
            continue

        part_name = getattr(image_part, "partname", None)
        filename = Path(str(part_name)).name if part_name is not None else None
        mime_type = getattr(image_part, "content_type", None) or _mime_type_for_filename(filename)
        image_id = _register_image_asset(
            assets,
            asset_lookup,
            data=image_part.blob,
            mime_type=mime_type,
            filename=filename,
        )

        image_counter += 1
        drawing_parent = element
        while drawing_parent is not None and not (
            isinstance(getattr(drawing_parent, "tag", None), str)
            and drawing_parent.tag.rsplit("}", 1)[-1] in {"inline", "anchor"}
        ):
            drawing_parent = drawing_parent.getparent()

        extent_el = None if drawing_parent is None else next(
            (
                child
                for child in drawing_parent.iter()
                if isinstance(getattr(child, "tag", None), str)
                and child.tag.rsplit("}", 1)[-1] == "extent"
                and child.get("cx") is not None
            ),
            None,
        )
        doc_pr_el = None if drawing_parent is None else next(
            (
                child
                for child in drawing_parent.iter()
                if isinstance(getattr(child, "tag", None), str)
                and child.tag.rsplit("}", 1)[-1] == "docPr"
            ),
            None,
        )
        images.append(
            ImageIR(
                unit_id=f"{paragraph_id}.img{image_counter}",
                image_id=image_id,
                alt_text=doc_pr_el.get("descr") if doc_pr_el is not None else None,
                title=doc_pr_el.get("name") if doc_pr_el is not None else None,
                display_width_pt=_emu_to_pt(extent_el.get("cx")) if extent_el is not None else None,
                display_height_pt=_emu_to_pt(extent_el.get("cy")) if extent_el is not None else None,
            )
        )

    return images, image_counter


def _parse_docx_paragraph_content(
    paragraph,
    paragraph_id: str,
    *,
    skip_empty: bool,
    normalizer: Callable[[str], str],
    assets: dict[str, ImageAsset],
    asset_lookup: dict[tuple[str, str], str],
) -> tuple[list[RunIR], list[ImageIR], list[object]]:
    runs: list[RunIR] = []
    images: list[ImageIR] = []
    content: list[object] = []
    image_counter = 0

    if not paragraph.runs:
        text = paragraph.text or ""
        if text or not skip_empty:
            run = RunIR(
                unit_id=f"{paragraph_id}.r1",
                text=text,
                normalized_text=normalizer(text),
            )
            runs.append(run)
            content.append(run)
        return runs, images, content

    for run_index, run in enumerate(paragraph.runs, start=1):
        text = run.text or ""
        run_images, image_counter = _parse_docx_run_images(
            run,
            paragraph_id=paragraph_id,
            image_counter=image_counter,
            assets=assets,
            asset_lookup=asset_lookup,
        )
        if text or (not skip_empty and not run_images):
            run_ir = RunIR(
                unit_id=f"{paragraph_id}.r{run_index}",
                text=text,
                normalized_text=normalizer(text),
            )
            runs.append(run_ir)
            content.append(run_ir)
        images.extend(run_images)
        content.extend(run_images)

    return runs, images, content


def _parse_docx_table(
    table,
    table_id: str,
    *,
    include_tables: bool,
    skip_empty: bool,
    normalizer: Callable[[str], str],
    assets: dict[str, ImageAsset],
    asset_lookup: dict[tuple[str, str], str],
    CT_P,
    CT_Tbl,
    Paragraph,
    Table,
) -> TableIR:
    table_ir = TableIR(unit_id=table_id)

    for tr_idx, row in enumerate(table.rows, start=1):
        for tc_idx, cell in enumerate(row.cells, start=1):
            cell_ir = TableCellIR(
                unit_id=f"{table_id}.tr{tr_idx}.tc{tc_idx}",
                row_index=tr_idx,
                col_index=tc_idx,
            )
            table_ir.cells.append(cell_ir)

            cp_idx = 0
            current_paragraph: ParagraphIR | None = None
            nested_table_counter_by_paragraph: dict[str, int] = {}

            for block in _iter_blocks_from_element(
                cell,
                cell._tc,
                CT_P=CT_P,
                CT_Tbl=CT_Tbl,
                Paragraph=Paragraph,
                Table=Table,
            ):
                if isinstance(block, Paragraph):
                    cp_idx += 1
                    paragraph_id = f"{table_id}.tr{tr_idx}.tc{tc_idx}.p{cp_idx}"
                    runs, images, content = _parse_docx_paragraph_content(
                        block,
                        paragraph_id,
                        skip_empty=skip_empty,
                        normalizer=normalizer,
                        assets=assets,
                        asset_lookup=asset_lookup,
                    )
                    current_paragraph = ParagraphIR(
                        unit_id=paragraph_id,
                        content=content,
                    )
                    current_paragraph.recompute_text(normalizer=normalizer)
                    if current_paragraph.content or current_paragraph.text or not skip_empty:
                        cell_ir.paragraphs.append(current_paragraph)
                    continue

                if not include_tables or not isinstance(block, Table):
                    continue

                if current_paragraph is None:
                    cp_idx += 1
                    current_paragraph = ParagraphIR(
                        unit_id=f"{table_id}.tr{tr_idx}.tc{tc_idx}.p{cp_idx}",
                    )
                    cell_ir.paragraphs.append(current_paragraph)

                table_counter = nested_table_counter_by_paragraph.get(current_paragraph.unit_id, 0) + 1
                nested_table_counter_by_paragraph[current_paragraph.unit_id] = table_counter
                nested_table = _parse_docx_table(
                    block,
                    f"{current_paragraph.unit_id}.tbl{table_counter}",
                    include_tables=include_tables,
                    skip_empty=skip_empty,
                    normalizer=normalizer,
                    assets=assets,
                    asset_lookup=asset_lookup,
                    CT_P=CT_P,
                    CT_Tbl=CT_Tbl,
                    Paragraph=Paragraph,
                    Table=Table,
                )
                current_paragraph.append_content(nested_table)
                current_paragraph.recompute_text(normalizer=normalizer)

            cell_ir.recompute_text(normalizer=normalizer)

    return table_ir


def _build_docx_doc_ir(
    source: "DocxDocument | str | Path | bytes",
    *,
    include_tables: bool,
    skip_empty: bool,
    source_path: str | Path | None,
    metadata: dict[str, Any] | None,
    normalizer: Callable[[str], str] | None,
    doc_id: str | None,
    doc_cls: type[DocIR] | None,
    **doc_kwargs: Any,
) -> DocIR:
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    normalize = normalizer or normalize_text_default
    doc = _load_docx_source(source)
    paragraphs: list[ParagraphIR] = []
    assets: dict[str, ImageAsset] = {}
    asset_lookup: dict[tuple[str, str], str] = {}

    p_idx = 0
    table_counter = 0

    for block in _iter_blocks(
        doc,
        CT_P=CT_P,
        CT_Tbl=CT_Tbl,
        Paragraph=Paragraph,
        Table=Table,
    ):
        if isinstance(block, Paragraph):
            p_idx += 1
            paragraph_id = f"s1.p{p_idx}"
            runs, images, content = _parse_docx_paragraph_content(
                block,
                paragraph_id,
                skip_empty=skip_empty,
                normalizer=normalize,
                assets=assets,
                asset_lookup=asset_lookup,
            )
            paragraph_ir = ParagraphIR(
                unit_id=paragraph_id,
                content=content,
            )
            paragraph_ir.recompute_text(normalizer=normalize)
            if paragraph_ir.content or paragraph_ir.text or not skip_empty:
                paragraphs.append(paragraph_ir)
            continue

        if not include_tables or not isinstance(block, Table):
            continue

        table_counter += 1
        p_idx += 1
        paragraph_id = f"s1.p{p_idx}"
        table_ir = _parse_docx_table(
            block,
            f"{paragraph_id}.r1.tbl{table_counter}",
            include_tables=include_tables,
            skip_empty=skip_empty,
            normalizer=normalize,
            assets=assets,
            asset_lookup=asset_lookup,
            CT_P=CT_P,
            CT_Tbl=CT_Tbl,
            Paragraph=Paragraph,
            Table=Table,
        )
        paragraph_ir = ParagraphIR(
            unit_id=paragraph_id,
            content=[table_ir],
        )
        paragraph_ir.recompute_text(normalizer=normalize)
        paragraphs.append(paragraph_ir)

    resolved_doc_id, resolved_source_path, resolved_metadata = _resolve_doc_metadata(
        source_path=source_path,
        source_doc_type="docx",
        metadata=metadata,
        doc_id=doc_id,
    )
    resolved_doc_cls = doc_cls or DocIR
    return resolved_doc_cls(
        doc_id=resolved_doc_id,
        source_path=resolved_source_path,
        source_doc_type="docx",
        metadata=resolved_metadata,
        assets=assets,
        paragraphs=paragraphs,
        **doc_kwargs,
    )


def _hwpx_binary_name_map(zf: zipfile.ZipFile) -> dict[str, str]:
    mapping: dict[str, str] = {}
    for name in zf.namelist():
        if not name.startswith("BinData/"):
            continue
        mapping[Path(name).stem.lower()] = name
    return mapping


def _find_hwpx_binary_path(binary_name_map: dict[str, str], binary_item_id: str | None) -> str | None:
    if not binary_item_id:
        return None
    return binary_name_map.get(binary_item_id.lower())


def _parse_hwpx_run_images(
    run_el: ET.Element,
    *,
    paragraph_id: str,
    image_counter: int,
    archive: zipfile.ZipFile,
    binary_name_map: dict[str, str],
    assets: dict[str, ImageAsset],
    asset_lookup: dict[tuple[str, str], str],
) -> tuple[list[ImageIR], int]:
    images: list[ImageIR] = []

    for child in list(run_el):
        tag = child.tag
        if tag != f"{_HP}pic":
            continue

        img_el = next(
            (
                element
                for element in child.iter()
                if isinstance(element.tag, str)
                and element.tag.rsplit("}", 1)[-1] == _HC_IMG_TAG
                and element.get("binaryItemIDRef")
            ),
            None,
        )
        if img_el is None:
            continue

        binary_path = _find_hwpx_binary_path(binary_name_map, img_el.get("binaryItemIDRef"))
        if binary_path is None:
            continue

        data = archive.read(binary_path)
        filename = Path(binary_path).name
        image_id = _register_image_asset(
            assets,
            asset_lookup,
            data=data,
            mime_type=_mime_type_for_filename(filename),
            filename=filename,
        )
        dim_el = child.find(f"{_HP}imgDim")
        display_width_pt = _hwpunit_to_pt(dim_el.get("dimwidth")) if dim_el is not None else None
        display_height_pt = _hwpunit_to_pt(dim_el.get("dimheight")) if dim_el is not None else None

        image_counter += 1
        images.append(
            ImageIR(
                unit_id=f"{paragraph_id}.img{image_counter}",
                image_id=image_id,
                display_width_pt=display_width_pt,
                display_height_pt=display_height_pt,
            )
        )

    return images, image_counter


def _parse_hwpx_paragraph_content(
    paragraph_el: ET.Element,
    paragraph_id: str,
    *,
    table_id_builder: Callable[[int], str],
    include_tables: bool,
    skip_empty: bool,
    normalizer: Callable[[str], str],
    archive: zipfile.ZipFile,
    binary_name_map: dict[str, str],
    assets: dict[str, ImageAsset],
    asset_lookup: dict[tuple[str, str], str],
) -> tuple[list[RunIR], list[ImageIR], list[TableIR], list[object]]:
    runs: list[RunIR] = []
    images: list[ImageIR] = []
    tables: list[TableIR] = []
    content: list[object] = []
    image_counter = 0
    table_counter = 0

    run_els = paragraph_el.findall(f"{_HP}run")
    if not run_els:
        text = _paragraph_text(paragraph_el)
        if text or not skip_empty:
            run = RunIR(
                unit_id=f"{paragraph_id}.r1",
                text=text,
                normalized_text=normalizer(text),
            )
            runs.append(run)
            content.append(run)
        return runs, images, tables, content

    for run_index, run_el in enumerate(run_els, start=1):
        text = _run_text(run_el)
        run_images, image_counter = _parse_hwpx_run_images(
            run_el,
            paragraph_id=paragraph_id,
            image_counter=image_counter,
            archive=archive,
            binary_name_map=binary_name_map,
            assets=assets,
            asset_lookup=asset_lookup,
        )
        run_tables: list[TableIR] = []
        if include_tables:
            for table_el in run_el.findall(f"{_HP}tbl"):
                table_counter += 1
                table_ir = _parse_hwpx_table(
                    table_el,
                    table_id_builder(table_counter),
                    include_tables=include_tables,
                    skip_empty=skip_empty,
                    normalizer=normalizer,
                    archive=archive,
                    binary_name_map=binary_name_map,
                    assets=assets,
                    asset_lookup=asset_lookup,
                )
                run_tables.append(table_ir)

        if text or (not skip_empty and not run_images and not run_tables):
            run = RunIR(
                unit_id=f"{paragraph_id}.r{run_index}",
                text=text,
                normalized_text=normalizer(text),
            )
            runs.append(run)
            content.append(run)
        images.extend(run_images)
        content.extend(run_images)
        tables.extend(run_tables)
        content.extend(run_tables)

    return runs, images, tables, content


def _parse_hwpx_cell_paragraphs(cell_el: ET.Element) -> list[ET.Element]:
    sub_list = cell_el.find(f"{_HP}subList")
    if sub_list is None:
        return []
    return [child for child in list(sub_list) if child.tag == f"{_HP}p"]


def _parse_hwpx_table(
    table_el: ET.Element,
    table_id: str,
    *,
    include_tables: bool,
    skip_empty: bool,
    normalizer: Callable[[str], str],
    archive: zipfile.ZipFile,
    binary_name_map: dict[str, str],
    assets: dict[str, ImageAsset],
    asset_lookup: dict[tuple[str, str], str],
) -> TableIR:
    table_ir = TableIR(unit_id=table_id)

    for tr_idx, row_el in enumerate(table_el.findall(f"{_HP}tr"), start=1):
        for tc_idx, cell_el in _logical_table_cells(row_el):
            cell_ir = TableCellIR(
                unit_id=f"{table_id}.tr{tr_idx}.tc{tc_idx}",
                row_index=tr_idx,
                col_index=tc_idx,
            )
            table_ir.cells.append(cell_ir)

            cell_paragraphs = _parse_hwpx_cell_paragraphs(cell_el)
            if not cell_paragraphs:
                if not skip_empty:
                    paragraph_ir = ParagraphIR(
                        unit_id=f"{table_id}.tr{tr_idx}.tc{tc_idx}.p1",
                        content=[
                            RunIR(
                                unit_id=f"{table_id}.tr{tr_idx}.tc{tc_idx}.p1.r1",
                                text="",
                                normalized_text=normalizer(""),
                            )
                        ],
                    )
                    paragraph_ir.recompute_text(normalizer=normalizer)
                    cell_ir.paragraphs.append(paragraph_ir)
                cell_ir.recompute_text(normalizer=normalizer)
                continue

            for cp_idx, paragraph_el in enumerate(cell_paragraphs, start=1):
                paragraph_id = f"{table_id}.tr{tr_idx}.tc{tc_idx}.p{cp_idx}"
                runs, images, tables, content = _parse_hwpx_paragraph_content(
                    paragraph_el,
                    paragraph_id,
                    table_id_builder=lambda counter, base=paragraph_id: f"{base}.tbl{counter}",
                    include_tables=include_tables,
                    skip_empty=skip_empty,
                    normalizer=normalizer,
                    archive=archive,
                    binary_name_map=binary_name_map,
                    assets=assets,
                    asset_lookup=asset_lookup,
                )
                paragraph_ir = ParagraphIR(
                    unit_id=paragraph_id,
                    content=content,
                )
                paragraph_ir.recompute_text(normalizer=normalizer)
                if paragraph_ir.content or paragraph_ir.text or not skip_empty:
                    cell_ir.paragraphs.append(paragraph_ir)

            cell_ir.recompute_text(normalizer=normalizer)

    return table_ir


def _build_hwpx_doc_ir(
    source: "HwpxDocument | str | Path | bytes",
    *,
    include_tables: bool,
    skip_empty: bool,
    source_path: str | Path | None,
    metadata: dict[str, Any] | None,
    normalizer: Callable[[str], str] | None,
    doc_id: str | None,
    doc_cls: type[DocIR] | None,
    **doc_kwargs: Any,
) -> DocIR:
    normalize = normalizer or normalize_text_default

    if isinstance(source, bytes):
        hwpx_bytes = source
    elif isinstance(source, (str, Path)):
        hwpx_bytes = Path(source).read_bytes()
    else:
        hwpx_bytes = coerce_source_to_supported_value(source, doc_type="hwpx")  # type: ignore[arg-type]
        if not isinstance(hwpx_bytes, bytes):
            raise TypeError(f"Unsupported HWPX source type: {type(source)!r}")

    section_roots = _section_roots_from_bytes(hwpx_bytes)
    assets: dict[str, ImageAsset] = {}
    asset_lookup: dict[tuple[str, str], str] = {}
    paragraphs: list[ParagraphIR] = []

    with zipfile.ZipFile(BytesIO(hwpx_bytes)) as archive:
        binary_name_map = _hwpx_binary_name_map(archive)

        for s_idx, section_root in enumerate(section_roots, start=1):
            for p_idx, paragraph_el in enumerate(section_root.findall(f"{_HP}p"), start=1):
                paragraph_id = f"s{s_idx}.p{p_idx}"
                runs, images, tables, content = _parse_hwpx_paragraph_content(
                    paragraph_el,
                    paragraph_id,
                    table_id_builder=lambda counter, base=paragraph_id: f"{base}.r1.tbl{counter}",
                    include_tables=include_tables,
                    skip_empty=skip_empty,
                    normalizer=normalize,
                    archive=archive,
                    binary_name_map=binary_name_map,
                    assets=assets,
                    asset_lookup=asset_lookup,
                )

                paragraph_ir = ParagraphIR(
                    unit_id=paragraph_id,
                    content=content,
                )
                paragraph_ir.recompute_text(normalizer=normalize)
                if paragraph_ir.content or paragraph_ir.text or not skip_empty:
                    paragraphs.append(paragraph_ir)

    resolved_doc_id, resolved_source_path, resolved_metadata = _resolve_doc_metadata(
        source_path=source_path,
        source_doc_type="hwpx",
        metadata=metadata,
        doc_id=doc_id,
    )
    resolved_doc_cls = doc_cls or DocIR
    return resolved_doc_cls(
        doc_id=resolved_doc_id,
        source_path=resolved_source_path,
        source_doc_type="hwpx",
        metadata=resolved_metadata,
        assets=assets,
        paragraphs=paragraphs,
        **doc_kwargs,
    )


def build_doc_ir_from_file(
    source: "HwpxDocument | DocxDocument | str | Path | bytes",
    *,
    doc_type: DocType,
    include_tables: bool = True,
    skip_empty: bool = False,
    source_path: str | Path | None = None,
    metadata: dict[str, Any] | None = None,
    normalizer: Callable[[str], str] | None = None,
    doc_id: str | None = None,
    doc_cls: type[DocIR] | None = None,
    **doc_kwargs: Any,
) -> DocIR:
    """Build document IR directly from a document source."""
    if doc_type == "docx":
        return _build_docx_doc_ir(
            source,
            include_tables=include_tables,
            skip_empty=skip_empty,
            source_path=source_path,
            metadata=metadata,
            normalizer=normalizer,
            doc_id=doc_id,
            doc_cls=doc_cls,
            **doc_kwargs,
        )

    if doc_type == "hwp":
        if not isinstance(source, (str, Path)):
            raise TypeError("HWP conversion currently requires a filesystem path.")
        return _build_hwpx_doc_ir(
            convert_hwp_to_hwpx_bytes(source),
            include_tables=include_tables,
            skip_empty=skip_empty,
            source_path=source_path,
            metadata=metadata,
            normalizer=normalizer,
            doc_id=doc_id,
            doc_cls=doc_cls,
            **doc_kwargs,
        )

    return _build_hwpx_doc_ir(
        source,
        include_tables=include_tables,
        skip_empty=skip_empty,
        source_path=source_path,
        metadata=metadata,
        normalizer=normalizer,
        doc_id=doc_id,
        doc_cls=doc_cls,
        **doc_kwargs,
    )


__all__ = ["build_doc_ir_from_file"]
