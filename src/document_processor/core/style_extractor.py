"""Unified style extraction for HWP/HWPX/DOCX documents."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
import re
from typing import TYPE_CHECKING, Literal
from xml.etree import ElementTree as ET
import zipfile

from ..io_utils import infer_doc_type
from ..style_types import CellStyleInfo, ParaStyleInfo, RunStyleInfo, StyleMap, TableStyleInfo
from .hwp_converter import convert_hwp_to_hwpx_bytes

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
    from ..hwpx import HwpxDocument


DocType = Literal["auto", "hwp", "hwpx", "docx", "pdf"]

_NS_HH = "http://www.hancom.co.kr/hwpml/2011/head"
_NS_HC = "http://www.hancom.co.kr/hwpml/2011/core"
_NS_HP = "http://www.hancom.co.kr/hwpml/2011/paragraph"
_HP = f"{{{_NS_HP}}}"

_HWPUNIT_PER_PT = 100.0
_HWPX_INHERIT_UINT = "4294967295"

_HWPX_BORDER_STYLE = {
    "SOLID": "solid",
    "DASH": "dashed",
    "DOT": "dotted",
    "DASH_DOT": "dashed",
    "DASH_DOT_DOT": "dotted",
    "DOUBLE": "double",
    "NONE": "none",
}

_HWPX_HALIGN = {
    "LEFT": "left",
    "CENTER": "center",
    "RIGHT": "right",
    "JUSTIFY": "justify",
    "DISTRIBUTE": "justify",
}

_HWPX_VALIGN = {
    "TOP": "top",
    "CENTER": "center",
    "BOTTOM": "bottom",
    "BASELINE": "top",
}

_DOCX_ALIGN = {0: "left", 1: "center", 2: "right", 3: "justify"}

_HWPX_VISIBLE_LINE_SHAPES = {
    "SOLID",
    "DASH",
    "DOT",
    "DASH_DOT",
    "DASH_DOT_DOT",
    "LONG_DASH",
    "CIRCLE",
    "DOUBLE",
    "SLIM_THICK",
    "THICK_SLIM",
    "SLIM_THICK_SLIM",
}


def _has_para_style(info: ParaStyleInfo) -> bool:
    return any(
        value is not None
        for value in (
            info.align,
            info.left_indent_pt,
            info.right_indent_pt,
            info.first_line_indent_pt,
            info.hanging_indent_pt,
            info.column_layout,
        )
    )


def _safe_int(value: str | None) -> int | None:
    if value is None:
        return None
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _length_to_pt(value) -> float | None:
    if value is None:
        return None
    try:
        return float(value) / 12700.0
    except (TypeError, ValueError):
        return None


def _docx_measure_to_pt(raw_value: str | None, measurement_type: str | None) -> float | None:
    if raw_value is None:
        return None
    try:
        number = float(raw_value)
    except (TypeError, ValueError):
        return None

    normalized_type = (measurement_type or "dxa").lower()
    if normalized_type == "dxa":
        return number / 20.0
    if normalized_type == "nil":
        return 0.0
    return None


def _docx_cell_margin_to_padding(margin_el) -> dict[str, float]:
    from docx.oxml.ns import qn

    padding: dict[str, float] = {}
    if margin_el is None:
        return padding

    for side in ("top", "right", "bottom", "left"):
        side_el = margin_el.find(qn(f"w:{side}"))
        if side_el is None:
            continue
        value_pt = _docx_measure_to_pt(
            side_el.get(qn("w:w")),
            side_el.get(qn("w:type")),
        )
        if value_pt is not None:
            padding[side] = value_pt

    return padding


def _apply_cell_padding(info: CellStyleInfo, padding: dict[str, float]) -> None:
    if "top" in padding:
        info.padding_top_pt = padding["top"]
    if "right" in padding:
        info.padding_right_pt = padding["right"]
    if "bottom" in padding:
        info.padding_bottom_pt = padding["bottom"]
    if "left" in padding:
        info.padding_left_pt = padding["left"]


def _hwp_numeric_to_pt(raw_value: str | None) -> float | None:
    if raw_value is None:
        return None
    raw_text = str(raw_value).strip()
    if raw_text in {_HWPX_INHERIT_UINT, "-1"}:
        return None
    try:
        return float(raw_text) / _HWPUNIT_PER_PT
    except (TypeError, ValueError):
        return None


def _hwp_margin_value_to_pt(el: ET.Element | None) -> float | None:
    if el is None:
        return None

    raw = el.get("value")
    if raw is None:
        return None
    raw_text = str(raw).strip()
    if raw_text == _HWPX_INHERIT_UINT:
        return None

    try:
        number = float(raw_text)
    except (TypeError, ValueError):
        return None

    unit = (el.get("unit") or "HWPUNIT").upper()
    if unit == "HWPUNIT":
        return number / _HWPUNIT_PER_PT
    if unit == "PT":
        return number
    return number


def _hwpx_table_cell_padding_defaults(table_el: ET.Element) -> dict[str, float]:
    margin_el = table_el.find(f"{_HP}inMargin")
    if margin_el is None:
        return {}

    padding: dict[str, float] = {}
    for side in ("top", "right", "bottom", "left"):
        value_pt = _hwp_numeric_to_pt(margin_el.get(side))
        if value_pt is not None:
            padding[side] = value_pt
    return padding


def _apply_hwpx_cell_margin(
    info: CellStyleInfo,
    margin_el: ET.Element | None,
    *,
    defaults: dict[str, float] | None = None,
) -> None:
    padding = dict(defaults or {})
    if margin_el is None and not padding:
        return

    if margin_el is not None:
        for side in ("top", "right", "bottom", "left"):
            value_pt = _hwp_numeric_to_pt(margin_el.get(side))
            if value_pt is not None:
                padding[side] = value_pt

    _apply_cell_padding(info, padding)


def _hwpx_border_css(border_el: ET.Element | None) -> str | None:
    if border_el is None:
        return None
    btype = border_el.get("type", "NONE")
    if btype == "NONE":
        return None
    width = border_el.get("width", "0.12 mm")
    try:
        mm_val = float(width.replace("mm", "").strip())
        px = max(1, round(mm_val * 3.78))
    except (ValueError, AttributeError):
        px = 1
    color = border_el.get("color", "#000000")
    style = _HWPX_BORDER_STYLE.get(btype, "solid")
    return f"{px}px {style} {color}"


def _hwpx_diagonal_border_css(
    border_fill_el: ET.Element,
    *,
    direction: Literal["slash", "backslash"],
) -> str | None:
    if direction == "slash":
        direction_el = border_fill_el.find(f"{{{_NS_HH}}}slash")
    else:
        direction_el = border_fill_el.find(f"{{{_NS_HH}}}backSlash")

    if direction_el is None or direction_el.get("type", "NONE") == "NONE":
        return None

    diagonal_el = border_fill_el.find(f"{{{_NS_HH}}}diagonal")
    return _hwpx_border_css(diagonal_el)


def _map_by_id(root: ET.Element | None, tag: str) -> dict[str, ET.Element]:
    if root is None:
        return {}
    out: dict[str, ET.Element] = {}
    for el in root.findall(f".//{{{_NS_HH}}}{tag}"):
        el_id = el.get("id")
        if el_id:
            out[el_id] = el
    return out


def _hwpx_para_style_from_pr(para_pr_el: ET.Element | None) -> ParaStyleInfo | None:
    if para_pr_el is None:
        return None

    info = ParaStyleInfo()

    align_el = para_pr_el.find(f"{{{_NS_HH}}}align")
    if align_el is not None:
        info.align = _HWPX_HALIGN.get(align_el.get("horizontal", ""))

    margin_el = para_pr_el.find(f"{{{_NS_HH}}}margin")
    if margin_el is None:
        margin_el = para_pr_el.find(f".//{{{_NS_HH}}}margin")
    if margin_el is not None:
        first_line = _hwp_margin_value_to_pt(margin_el.find(f"{{{_NS_HC}}}intent"))
        info.first_line_indent_pt = first_line
        info.left_indent_pt = _hwp_margin_value_to_pt(margin_el.find(f"{{{_NS_HC}}}left"))
        info.right_indent_pt = _hwp_margin_value_to_pt(margin_el.find(f"{{{_NS_HC}}}right"))
        if first_line is not None and first_line < 0:
            info.hanging_indent_pt = abs(first_line)

    return info if _has_para_style(info) else None


def _hwpx_run_style_from_char_pr(char_pr_el: ET.Element | None) -> RunStyleInfo:
    info = RunStyleInfo()
    if char_pr_el is None:
        return info

    info.bold = char_pr_el.find(f"{{{_NS_HH}}}bold") is not None
    info.italic = char_pr_el.find(f"{{{_NS_HH}}}italic") is not None

    underline_el = char_pr_el.find(f"{{{_NS_HH}}}underline")
    if underline_el is not None and underline_el.get("type", "NONE") != "NONE":
        info.underline = True

    strike_el = char_pr_el.find(f"{{{_NS_HH}}}strikeout")
    if strike_el is not None:
        strike_type = (strike_el.get("type") or "").upper()
        strike_shape = (strike_el.get("shape") or "").upper()
        if strike_type and strike_type != "NONE":
            info.strikethrough = True
        elif strike_shape in _HWPX_VISIBLE_LINE_SHAPES:
            info.strikethrough = True

    color = char_pr_el.get("textColor")
    if color and color != "#000000":
        info.color = color

    height = char_pr_el.get("height")
    if height is not None:
        try:
            info.size_pt = int(height) / 100.0
        except (TypeError, ValueError):
            pass

    return info


def _iter_section_paragraphs(section_root: ET.Element) -> list[ET.Element]:
    return section_root.findall(f"{_HP}p")


def _iter_paragraph_tables(paragraph_el: ET.Element) -> list[ET.Element]:
    return paragraph_el.findall(f"{_HP}run/{_HP}tbl")


def _iter_cell_paragraphs(cell_el: ET.Element) -> list[ET.Element]:
    direct = cell_el.findall(f"{_HP}subList/{_HP}p")
    if direct:
        return direct
    return cell_el.findall(f".//{_HP}p")


def _logical_table_cells(row_el: ET.Element) -> list[tuple[int, ET.Element]]:
    """Return logical 1-based column indices for row cells."""
    logical_cells: list[tuple[int, ET.Element]] = []
    fallback_col = 1

    for cell_el in row_el.findall(f"{_HP}tc"):
        cell_addr = cell_el.find(f"{_HP}cellAddr")
        col_addr = _safe_int(cell_addr.get("colAddr")) if cell_addr is not None else None
        logical_col = (col_addr + 1) if col_addr is not None else fallback_col
        logical_cells.append((logical_col, cell_el))

        cell_span = cell_el.find(f"{_HP}cellSpan")
        colspan = _safe_int(cell_span.get("colSpan")) if cell_span is not None else None
        fallback_col = max(fallback_col, logical_col + max(colspan or 1, 1))

    return logical_cells


def _hwpx_table_dimensions(table_el: ET.Element) -> tuple[int, int]:
    row_els = table_el.findall(f"{_HP}tr")
    table_row_count = 0
    table_col_count = 0
    for row_el in row_els:
        for logical_col, cell_el in _logical_table_cells(row_el):
            cell_addr = cell_el.find(f"{_HP}cellAddr")
            row_addr = _safe_int(cell_addr.get("rowAddr")) if cell_addr is not None else None
            logical_row = (row_addr + 1) if row_addr is not None else 1
            cell_span = cell_el.find(f"{_HP}cellSpan")
            rowspan = _safe_int(cell_span.get("rowSpan")) if cell_span is not None else None
            colspan = _safe_int(cell_span.get("colSpan")) if cell_span is not None else None
            table_row_count = max(table_row_count, logical_row + max(rowspan or 1, 1) - 1)
            table_col_count = max(table_col_count, logical_col + max(colspan or 1, 1) - 1)
    return table_row_count or len(row_els), table_col_count


def _hwpx_table_size(table_el: ET.Element) -> tuple[float | None, float | None]:
    size_el = table_el.find(f"{_HP}sz")
    if size_el is None:
        return None, None
    return (
        _hwp_numeric_to_pt(size_el.get("width")),
        _hwp_numeric_to_pt(size_el.get("height")),
    )


def _extract_hwpx_table_styles(
    style_map: StyleMap,
    table_el: ET.Element,
    table_id: str,
    *,
    para_pr_map: dict[str, ET.Element],
    char_pr_map: dict[str, ET.Element],
    border_fill_map: dict[str, ET.Element],
) -> None:
    row_count, col_count = _hwpx_table_dimensions(table_el)
    width_pt, height_pt = _hwpx_table_size(table_el)
    table_cell_padding_defaults = _hwpx_table_cell_padding_defaults(table_el)
    style_map.tables[table_id] = TableStyleInfo(
        row_count=row_count,
        col_count=col_count,
        width_pt=width_pt,
        height_pt=height_pt,
    )

    for tr_idx, row_el in enumerate(table_el.findall(f"{_HP}tr"), start=1):
        for tc_idx, cell_el in _logical_table_cells(row_el):
            cell_id = f"{table_id}.tr{tr_idx}.tc{tc_idx}"
            style_map.cells[cell_id] = _hwpx_cell_style(
                cell_el,
                para_pr_map=para_pr_map,
                border_fill_map=border_fill_map,
                table_cell_padding_defaults=table_cell_padding_defaults,
            )

            cell_paragraphs = _iter_cell_paragraphs(cell_el)
            if not cell_paragraphs:
                style_map.runs[f"{cell_id}.p1.r1"] = RunStyleInfo()
                continue

            for cp_idx, cell_para_el in enumerate(cell_paragraphs, start=1):
                cell_paragraph_id = f"{cell_id}.p{cp_idx}"
                cell_para_pr_ref = cell_para_el.get("paraPrIDRef")
                if cell_para_pr_ref and cell_para_pr_ref in para_pr_map:
                    cp_style = _hwpx_para_style_from_pr(para_pr_map[cell_para_pr_ref])
                    if cp_style is not None:
                        style_map.paragraphs[cell_paragraph_id] = cp_style

                cell_run_els = cell_para_el.findall(f"{_HP}run")
                if not cell_run_els:
                    style_map.runs[f"{cell_paragraph_id}.r1"] = RunStyleInfo()
                else:
                    for cr_idx, cell_run_el in enumerate(cell_run_els, start=1):
                        char_pr_ref = cell_run_el.get("charPrIDRef")
                        char_pr_el = char_pr_map.get(char_pr_ref) if char_pr_ref else None
                        style_map.runs[f"{cell_paragraph_id}.r{cr_idx}"] = (
                            _hwpx_run_style_from_char_pr(char_pr_el)
                        )

                for nested_t_idx, nested_table_el in enumerate(_iter_paragraph_tables(cell_para_el), start=1):
                    nested_table_id = f"{cell_paragraph_id}.tbl{nested_t_idx}"
                    _extract_hwpx_table_styles(
                        style_map,
                        nested_table_el,
                        nested_table_id,
                        para_pr_map=para_pr_map,
                        char_pr_map=char_pr_map,
                        border_fill_map=border_fill_map,
                    )


def _section_roots_from_bytes(source: bytes) -> list[ET.Element]:
    section_name_pattern = re.compile(r"^Contents/section\d+\.xml$")

    with zipfile.ZipFile(BytesIO(source)) as zf:
        def _section_order(name: str) -> int:
            match = re.search(r"section(\d+)\.xml$", name)
            return int(match.group(1)) if match else -1

        names = sorted(
            (name for name in zf.namelist() if section_name_pattern.match(name)),
            key=_section_order,
        )
        return [ET.fromstring(zf.read(name)) for name in names]


def _header_root_from_bytes(source: bytes) -> ET.Element | None:
    with zipfile.ZipFile(BytesIO(source)) as zf:
        try:
            return ET.fromstring(zf.read("Contents/header.xml"))
        except KeyError:
            return None


def _hwpx_cell_style(
    cell_el: ET.Element,
    *,
    para_pr_map: dict[str, ET.Element],
    border_fill_map: dict[str, ET.Element],
    table_cell_padding_defaults: dict[str, float] | None = None,
) -> CellStyleInfo:
    info = CellStyleInfo()

    span_el = cell_el.find(f"{_HP}cellSpan")
    if span_el is not None:
        info.rowspan = _safe_int(span_el.get("rowSpan")) or 1
        info.colspan = _safe_int(span_el.get("colSpan")) or 1

    sub_list = cell_el.find(f"{_HP}subList")
    if sub_list is not None:
        valign = sub_list.get("vertAlign", "")
        info.vertical_align = _HWPX_VALIGN.get(valign)

    cell_paragraphs = _iter_cell_paragraphs(cell_el)
    if cell_paragraphs:
        first_para = cell_paragraphs[0]
        pp_ref = first_para.get("paraPrIDRef")
        if pp_ref and pp_ref in para_pr_map:
            pstyle = _hwpx_para_style_from_pr(para_pr_map[pp_ref])
            if pstyle and pstyle.align:
                info.horizontal_align = pstyle.align

    bf_ref = cell_el.get("borderFillIDRef")
    if bf_ref and bf_ref in border_fill_map:
        border_fill = border_fill_map[bf_ref]
        info.border_top = _hwpx_border_css(border_fill.find(f"{{{_NS_HH}}}topBorder"))
        info.border_bottom = _hwpx_border_css(border_fill.find(f"{{{_NS_HH}}}bottomBorder"))
        info.border_left = _hwpx_border_css(border_fill.find(f"{{{_NS_HH}}}leftBorder"))
        info.border_right = _hwpx_border_css(border_fill.find(f"{{{_NS_HH}}}rightBorder"))
        info.diagonal_tr_bl = _hwpx_diagonal_border_css(border_fill, direction="slash")
        info.diagonal_tl_br = _hwpx_diagonal_border_css(border_fill, direction="backslash")

        fill_brush = border_fill.find(f"{{{_NS_HH}}}fillBrush")
        if fill_brush is None:
            fill_brush = border_fill.find(f"{{{_NS_HC}}}fillBrush")
        if fill_brush is not None:
            face_color = fill_brush.get("faceColor")
            if not face_color:
                win_brush = fill_brush.find(f"{{{_NS_HH}}}winBrush")
                if win_brush is None:
                    win_brush = fill_brush.find(f"{{{_NS_HC}}}winBrush")
                if win_brush is not None:
                    face_color = win_brush.get("faceColor")
            if face_color and face_color.lower() not in ("none", "#ffffff", "transparent"):
                info.background = face_color

    cell_size = cell_el.find(f"{_HP}cellSz")
    if cell_size is not None:
        info.width_pt = _hwp_numeric_to_pt(cell_size.get("width"))
        info.height_pt = _hwp_numeric_to_pt(cell_size.get("height"))

    _apply_hwpx_cell_margin(
        info,
        cell_el.find(f"{_HP}cellMargin"),
        defaults=table_cell_padding_defaults,
    )

    return info


def _extract_styles_hwpx_from_roots(
    section_roots: list[ET.Element],
    *,
    header_root: ET.Element | None,
) -> StyleMap:
    style_map = StyleMap()

    para_pr_map = _map_by_id(header_root, "paraPr")
    char_pr_map = _map_by_id(header_root, "charPr")
    border_fill_map = _map_by_id(header_root, "borderFill")

    for s_idx, section_root in enumerate(section_roots, start=1):
        for p_idx, para_el in enumerate(_iter_section_paragraphs(section_root), start=1):
            paragraph_id = f"s{s_idx}.p{p_idx}"

            para_pr_ref = para_el.get("paraPrIDRef")
            if para_pr_ref and para_pr_ref in para_pr_map:
                para_style = _hwpx_para_style_from_pr(para_pr_map[para_pr_ref])
                if para_style is not None:
                    style_map.paragraphs[paragraph_id] = para_style

            run_els = para_el.findall(f"{_HP}run")
            if not run_els:
                style_map.runs[f"{paragraph_id}.r1"] = RunStyleInfo()
            else:
                for r_idx, run_el in enumerate(run_els, start=1):
                    char_pr_ref = run_el.get("charPrIDRef")
                    char_pr_el = char_pr_map.get(char_pr_ref) if char_pr_ref else None
                    style_map.runs[f"{paragraph_id}.r{r_idx}"] = _hwpx_run_style_from_char_pr(
                        char_pr_el
                    )

            for t_idx, table_el in enumerate(_iter_paragraph_tables(para_el), start=1):
                table_id = f"{paragraph_id}.r1.tbl{t_idx}"
                _extract_hwpx_table_styles(
                    style_map,
                    table_el,
                    table_id,
                    para_pr_map=para_pr_map,
                    char_pr_map=char_pr_map,
                    border_fill_map=border_fill_map,
                )

    return style_map


def extract_styles_hwpx(source: "HwpxDocument | str | Path | bytes") -> StyleMap:
    """Extract style map from HWPX source."""
    from ..hwpx import HwpxDocument

    if isinstance(source, bytes):
        return _extract_styles_hwpx_from_roots(
            _section_roots_from_bytes(source),
            header_root=_header_root_from_bytes(source),
        )

    if isinstance(source, (str, Path)):
        return extract_styles_hwpx(Path(source).read_bytes())

    if isinstance(source, HwpxDocument):
        section_roots = [section.element for section in source.sections]
        header_root = source.headers[0].element if source.headers else None
        return _extract_styles_hwpx_from_roots(section_roots, header_root=header_root)

    raise TypeError(
        "source must be HwpxDocument, bytes, or a .hwpx path, "
        f"got {type(source)!r}"
    )


def _docx_run_style(run) -> RunStyleInfo:
    info = RunStyleInfo(
        bold=bool(run.bold),
        italic=bool(run.italic),
        underline=bool(run.underline),
    )

    font = run.font

    if font.color and font.color.rgb:
        rgb = str(font.color.rgb)
        if rgb != "000000":
            info.color = f"#{rgb}"

    if font.size is not None:
        info.size_pt = _length_to_pt(font.size)

    if font.strike:
        info.strikethrough = True
    if font.superscript:
        info.superscript = True
    if font.subscript:
        info.subscript = True
    if font.highlight_color is not None:
        info.highlight = str(font.highlight_color)

    return info


def _docx_para_style(paragraph) -> ParaStyleInfo | None:
    info = ParaStyleInfo()

    if paragraph.alignment is not None:
        info.align = _DOCX_ALIGN.get(paragraph.alignment)

    pf = paragraph.paragraph_format
    if pf is not None:
        info.left_indent_pt = _length_to_pt(pf.left_indent)
        info.right_indent_pt = _length_to_pt(pf.right_indent)

        first_line = _length_to_pt(pf.first_line_indent)
        info.first_line_indent_pt = first_line
        if first_line is not None and first_line < 0:
            info.hanging_indent_pt = abs(first_line)

    return info if _has_para_style(info) else None


def _docx_border_css(tc_borders, side: str) -> str | None:
    if tc_borders is None:
        return None

    from docx.oxml.ns import qn

    el = tc_borders.find(qn(f"w:{side}"))
    if el is None:
        return None
    val = el.get(qn("w:val"), "none")
    if val in ("none", "nil"):
        return None
    sz = el.get(qn("w:sz"), "4")
    try:
        px = max(1, round(int(sz) / 8 * 1.333))
    except (TypeError, ValueError):
        px = 1
    color = el.get(qn("w:color"), "000000")
    if color.lower() == "auto":
        color = "000000"
    style_map = {"single": "solid", "double": "double", "dashed": "dashed", "dotted": "dotted"}
    css_style = style_map.get(val, "solid")
    return f"{px}px {css_style} #{color}"


def _docx_table_style_border_defaults(
    style_id: str | None,
    *,
    style_elements: dict[str, object],
    cache: dict[str, dict[str, str | None]],
) -> dict[str, str | None]:
    from docx.oxml.ns import qn

    if not style_id:
        return {}
    cached = cache.get(style_id)
    if cached is not None:
        return dict(cached)

    style_el = style_elements.get(style_id)
    if style_el is None:
        cache[style_id] = {}
        return {}

    based_on_el = style_el.find(qn("w:basedOn"))
    base_style_id = based_on_el.get(qn("w:val")) if based_on_el is not None else None
    merged = _docx_table_style_border_defaults(
        base_style_id,
        style_elements=style_elements,
        cache=cache,
    )

    tbl_pr = style_el.find(qn("w:tblPr"))
    tbl_borders = tbl_pr.find(qn("w:tblBorders")) if tbl_pr is not None else None
    if tbl_borders is not None:
        for border_name, border_key in (
            ("top", "top"),
            ("bottom", "bottom"),
            ("left", "left"),
            ("right", "right"),
            ("insideH", "inside_h"),
            ("insideV", "inside_v"),
        ):
            border_css = _docx_border_css(tbl_borders, border_name)
            if border_css is not None:
                merged[border_key] = border_css

    cache[style_id] = dict(merged)
    return merged


def _docx_table_style_cell_padding_defaults(
    style_id: str | None,
    *,
    style_elements: dict[str, object],
    cache: dict[str, dict[str, float]],
) -> dict[str, float]:
    from docx.oxml.ns import qn

    if not style_id:
        return {}
    cached = cache.get(style_id)
    if cached is not None:
        return dict(cached)

    style_el = style_elements.get(style_id)
    if style_el is None:
        cache[style_id] = {}
        return {}

    based_on_el = style_el.find(qn("w:basedOn"))
    base_style_id = based_on_el.get(qn("w:val")) if based_on_el is not None else None
    merged = _docx_table_style_cell_padding_defaults(
        base_style_id,
        style_elements=style_elements,
        cache=cache,
    )

    tbl_pr = style_el.find(qn("w:tblPr"))
    cell_margin = tbl_pr.find(qn("w:tblCellMar")) if tbl_pr is not None else None
    merged.update(_docx_cell_margin_to_padding(cell_margin))

    cache[style_id] = dict(merged)
    return merged


def _docx_table_cell_padding_defaults(
    table,
    table_style_id: str | None,
    *,
    style_elements: dict[str, object],
    cache: dict[str, dict[str, float]],
) -> dict[str, float]:
    from docx.oxml.ns import qn

    defaults = _docx_table_style_cell_padding_defaults(
        table_style_id,
        style_elements=style_elements,
        cache=cache,
    )

    tbl_pr = table._tbl.find(qn("w:tblPr"))
    cell_margin = tbl_pr.find(qn("w:tblCellMar")) if tbl_pr is not None else None
    defaults.update(_docx_cell_margin_to_padding(cell_margin))
    return defaults


def _docx_default_cell_border(
    side: str,
    *,
    row_index: int,
    col_index: int,
    row_count: int,
    col_count: int,
    table_border_defaults: dict[str, str | None],
) -> str | None:
    if side == "top":
        return table_border_defaults.get("top") if row_index == 1 else table_border_defaults.get("inside_h")
    if side == "bottom":
        return table_border_defaults.get("bottom") if row_index == row_count else table_border_defaults.get("inside_h")
    if side == "left":
        return table_border_defaults.get("left") if col_index == 1 else table_border_defaults.get("inside_v")
    if side == "right":
        return table_border_defaults.get("right") if col_index == col_count else table_border_defaults.get("inside_v")
    return None


def _docx_table_size(table) -> tuple[float | None, float | None]:
    from docx.oxml.ns import qn

    table_width_pt: float | None = None
    table_height_pt: float | None = None

    tbl_pr = table._tbl.find(qn("w:tblPr"))
    if tbl_pr is not None:
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is not None:
            table_width_pt = _docx_measure_to_pt(
                tbl_w.get(qn("w:w")),
                tbl_w.get(qn("w:type")),
            )

    if table_width_pt is None:
        tbl_grid = table._tbl.find(qn("w:tblGrid"))
        if tbl_grid is not None:
            grid_width_pt = 0.0
            has_grid = False
            for grid_col in tbl_grid.findall(qn("w:gridCol")):
                width_pt = _docx_measure_to_pt(grid_col.get(qn("w:w")), "dxa")
                if width_pt is not None:
                    has_grid = True
                    grid_width_pt += width_pt
            if has_grid:
                table_width_pt = grid_width_pt

    row_height_total = 0.0
    has_row_height = False
    for row in table.rows:
        tr_pr = row._tr.find(qn("w:trPr"))
        tr_height = tr_pr.find(qn("w:trHeight")) if tr_pr is not None else None
        if tr_height is None:
            continue
        height_pt = _docx_measure_to_pt(tr_height.get(qn("w:val")), "dxa")
        if height_pt is None:
            continue
        has_row_height = True
        row_height_total += height_pt
    if has_row_height:
        table_height_pt = row_height_total

    return table_width_pt, table_height_pt


def _docx_cell_style(
    cell,
    *,
    row_index: int,
    col_index: int,
    row_count: int,
    col_count: int,
    row_height_pt: float | None = None,
    table_border_defaults: dict[str, str | None] | None = None,
    table_cell_padding_defaults: dict[str, float] | None = None,
) -> CellStyleInfo:
    from docx.oxml.ns import qn

    info = CellStyleInfo()
    tc = cell._tc
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        return info

    grid_span = tc_pr.find(qn("w:gridSpan"))
    if grid_span is not None:
        try:
            info.colspan = int(grid_span.get(qn("w:val"), "1"))
        except (TypeError, ValueError):
            pass

    v_align = tc_pr.find(qn("w:vAlign"))
    if v_align is not None:
        val = v_align.get(qn("w:val"), "")
        info.vertical_align = {"top": "top", "center": "center", "bottom": "bottom"}.get(val)

    if cell.paragraphs:
        pstyle = _docx_para_style(cell.paragraphs[0])
        if pstyle is not None and pstyle.align is not None:
            info.horizontal_align = pstyle.align

    tc_width = tc_pr.find(qn("w:tcW"))
    if tc_width is not None:
        info.width_pt = _docx_measure_to_pt(
            tc_width.get(qn("w:w")),
            tc_width.get(qn("w:type")),
        )
    if row_height_pt is not None:
        info.height_pt = row_height_pt

    table_cell_padding_defaults = table_cell_padding_defaults or {}
    if table_cell_padding_defaults:
        _apply_cell_padding(info, table_cell_padding_defaults)
    cell_margin = tc_pr.find(qn("w:tcMar"))
    if cell_margin is not None:
        _apply_cell_padding(info, _docx_cell_margin_to_padding(cell_margin))

    shd = tc_pr.find(qn("w:shd"))
    if shd is not None:
        fill = shd.get(qn("w:fill"))
        if fill and fill.lower() not in ("auto", "ffffff", "none"):
            info.background = f"#{fill}"

    tc_borders = tc_pr.find(qn("w:tcBorders"))
    table_border_defaults = table_border_defaults or {}
    info.border_top = _docx_border_css(tc_borders, "top") or _docx_default_cell_border(
        "top",
        row_index=row_index,
        col_index=col_index,
        row_count=row_count,
        col_count=col_count,
        table_border_defaults=table_border_defaults,
    )
    info.border_bottom = _docx_border_css(tc_borders, "bottom") or _docx_default_cell_border(
        "bottom",
        row_index=row_index,
        col_index=col_index,
        row_count=row_count,
        col_count=col_count,
        table_border_defaults=table_border_defaults,
    )
    info.border_left = _docx_border_css(tc_borders, "left") or _docx_default_cell_border(
        "left",
        row_index=row_index,
        col_index=col_index,
        row_count=row_count,
        col_count=col_count,
        table_border_defaults=table_border_defaults,
    )
    info.border_right = _docx_border_css(tc_borders, "right") or _docx_default_cell_border(
        "right",
        row_index=row_index,
        col_index=col_index,
        row_count=row_count,
        col_count=col_count,
        table_border_defaults=table_border_defaults,
    )
    info.diagonal_tl_br = _docx_border_css(tc_borders, "tl2br")
    info.diagonal_tr_bl = _docx_border_css(tc_borders, "tr2bl")
    return info


def extract_styles_docx(
    source: "DocxDocument | str | Path | bytes",
    *,
    include_tables: bool = True,
) -> StyleMap:
    """Extract style map from DOCX source."""
    from docx import Document as load_docx
    from docx.document import Document as DocxDocument
    from docx.oxml.ns import qn
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    from .docx_structured_exporter import _iter_blocks, _iter_blocks_from_element

    if isinstance(source, DocxDocument):
        doc = source
    elif isinstance(source, bytes):
        doc = load_docx(BytesIO(source))
    else:
        doc = load_docx(str(source))

    style_map = StyleMap()
    style_elements: dict[str, object] = {}
    border_defaults_cache: dict[str, dict[str, str | None]] = {}
    cell_padding_defaults_cache: dict[str, dict[str, float]] = {}

    for style_el in doc.styles.element.findall(qn("w:style")):
        if style_el.get(qn("w:type")) != "table":
            continue
        style_id = style_el.get(qn("w:styleId"))
        if style_id:
            style_elements[style_id] = style_el

    p_idx = 0
    tbl_counter = 0

    def _extract_docx_table_styles(table, table_id: str) -> None:
        vmerge_starts: dict[tuple[str, int], str] = {}
        table_style_id = table.style.style_id if table.style is not None else None
        table_border_defaults = _docx_table_style_border_defaults(
            table_style_id,
            style_elements=style_elements,
            cache=border_defaults_cache,
        )
        table_cell_padding_defaults = _docx_table_cell_padding_defaults(
            table,
            table_style_id,
            style_elements=style_elements,
            cache=cell_padding_defaults_cache,
        )
        table_width_pt, table_height_pt = _docx_table_size(table)

        style_map.tables[table_id] = TableStyleInfo(
            row_count=len(table.rows),
            col_count=len(table.columns),
            width_pt=table_width_pt,
            height_pt=table_height_pt,
        )

        for tr_idx, row in enumerate(table.rows, start=1):
            tr_pr = row._tr.find(qn("w:trPr"))
            tr_height = tr_pr.find(qn("w:trHeight")) if tr_pr is not None else None
            row_height_pt = None
            if tr_height is not None:
                row_height_pt = _docx_measure_to_pt(
                    tr_height.get(qn("w:val")),
                    "dxa",
                )
            for tc_idx, cell in enumerate(row.cells, start=1):
                cell_id = f"{table_id}.tr{tr_idx}.tc{tc_idx}"

                tc_pr = cell._tc.find(qn("w:tcPr"))
                if tc_pr is not None:
                    v_merge = tc_pr.find(qn("w:vMerge"))
                    if v_merge is not None:
                        val = v_merge.get(qn("w:val"), "")
                        col_key = (table_id, tc_idx)
                        if val == "restart":
                            vmerge_starts[col_key] = cell_id
                        elif col_key in vmerge_starts:
                            start_cell_id = vmerge_starts[col_key]
                            start_style = style_map.cells.get(start_cell_id)
                            if start_style is not None:
                                start_style.rowspan += 1
                            continue
                    else:
                        vmerge_starts.pop((table_id, tc_idx), None)

                style_map.cells[cell_id] = _docx_cell_style(
                    cell,
                    row_index=tr_idx,
                    col_index=tc_idx,
                    row_count=len(table.rows),
                    col_count=len(table.columns),
                    row_height_pt=row_height_pt,
                    table_border_defaults=table_border_defaults,
                    table_cell_padding_defaults=table_cell_padding_defaults,
                )

                cp_idx = 0
                current_paragraph_id: str | None = None
                nested_table_counter_by_paragraph: dict[str, int] = {}

                for block in _iter_blocks_from_element(
                    cell,
                    cell._tc,
                    CT_P=CT_P,
                    CT_Tbl=CT_Tbl,
                    Paragraph=Paragraph,
                    Table=Table,
                ):
                    if isinstance(block, Paragraph):
                        cp_idx += 1
                        current_paragraph_id = f"{cell_id}.p{cp_idx}"
                        cp_style = _docx_para_style(block)
                        if cp_style is not None:
                            style_map.paragraphs[current_paragraph_id] = cp_style

                        if not block.runs:
                            style_map.runs[f"{current_paragraph_id}.r1"] = RunStyleInfo()
                            continue

                        for cr_idx, cell_run in enumerate(block.runs, start=1):
                            style_map.runs[f"{current_paragraph_id}.r{cr_idx}"] = _docx_run_style(
                                cell_run
                            )
                        continue

                    if current_paragraph_id is None:
                        cp_idx += 1
                        current_paragraph_id = f"{cell_id}.p{cp_idx}"

                    nested_tbl_counter = nested_table_counter_by_paragraph.get(current_paragraph_id, 0) + 1
                    nested_table_counter_by_paragraph[current_paragraph_id] = nested_tbl_counter
                    nested_table_id = f"{current_paragraph_id}.tbl{nested_tbl_counter}"
                    _extract_docx_table_styles(block, nested_table_id)

    for block in _iter_blocks(
        doc,
        CT_P=CT_P,
        CT_Tbl=CT_Tbl,
        Paragraph=Paragraph,
        Table=Table,
    ):
        if isinstance(block, Paragraph):
            p_idx += 1
            paragraph_id = f"s1.p{p_idx}"

            pstyle = _docx_para_style(block)
            if pstyle is not None:
                style_map.paragraphs[paragraph_id] = pstyle

            if not block.runs:
                style_map.runs[f"{paragraph_id}.r1"] = RunStyleInfo()
            else:
                for r_idx, run in enumerate(block.runs, start=1):
                    style_map.runs[f"{paragraph_id}.r{r_idx}"] = _docx_run_style(run)
            continue

        if not include_tables or not isinstance(block, Table):
            continue

        tbl_counter += 1
        p_idx += 1
        table_id = f"s1.p{p_idx}.r1.tbl{tbl_counter}"
        _extract_docx_table_styles(block, table_id)

    return style_map


def extract_styles(
    source: "HwpxDocument | DocxDocument | str | Path | bytes",
    *,
    doc_type: DocType = "auto",
    include_tables: bool = True,
) -> StyleMap:
    """Extract styles for HWP/HWPX/DOCX with one interface."""
    resolved = infer_doc_type(source, doc_type)

    if resolved == "pdf":
        raise NotImplementedError("PDF style extraction is not implemented yet.")

    if resolved == "hwp":
        if not isinstance(source, (str, Path)):
            raise TypeError("HWP conversion currently requires a filesystem path.")
        hwpx_bytes = convert_hwp_to_hwpx_bytes(source)
        return extract_styles_hwpx(hwpx_bytes)

    if resolved == "hwpx":
        return extract_styles_hwpx(source)

    return extract_styles_docx(source, include_tables=include_tables)


__all__ = [
    "DocType",
    "extract_styles",
    "extract_styles_docx",
    "extract_styles_hwpx",
]
